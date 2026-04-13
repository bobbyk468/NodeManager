"""
Build ConceptGrade_System_Overview.docx

This module generates the complete ConceptGrade System Overview Word document.
It creates four matplotlib diagrams (architecture, KG example, brushing flow, pipeline)
as PNG images, then embeds them into a formatted Word document with:
  - Heading hierarchy (H1–H4) for automatic TOC generation
  - Captions with bold labels and italic descriptions
  - Tables with colour-coded headers and alternating row shading
  - Callout boxes for key insights
  - Proper spacing and alignment throughout

Key design principles:
  - Diagrams use 150 DPI for clarity in both screen and print
  - Node positions calculated to avoid overlaps (e.g. KG nodes spread to corners)
  - Arrows sized for visibility at Word document scale (lw=3.5, mutation_scale=28)
  - Colour palette aligned across all diagrams for visual consistency
"""

import io
import os
import matplotlib
matplotlib.use('Agg')  # non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Colour palette ────────────────────────────────────────────────────────────
BLUE   = '#1e40af'
TEAL   = '#0f766e'
AMBER  = '#b45309'
GREEN  = '#15803d'
RED    = '#b91c1c'
GREY   = '#374151'
LIGHT  = '#f0f9ff'
WHITE  = '#ffffff'
BG     = '#f8fafc'


# ═══════════════════════════════════════════════════════════════════════════════
# DIAGRAM HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def save_fig(fig, name: str) -> str:
    """Save figure to a PNG file next to this script and return the path."""
    path = os.path.join(OUT_DIR, name)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def diagram_architecture() -> str:
    """Three-tier architecture box diagram."""
    fig, ax = plt.subplots(figsize=(9, 5))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')

    layers = [
        (BLUE,  '#dbeafe', 'Layer 1 — Python Pipeline (concept-aware)',
         'Reads student answers  ·  Builds Knowledge Graphs\n'
         'Grades answers (5-stage KG pipeline)  ·  Saves results as JSON files',
         4.5),
        (TEAL,  '#ccfbf1', 'Layer 2 — NestJS Backend API (TypeScript)',
         'Reads JSON files  ·  Serves data via REST endpoints\n'
         'Handles study event logging  ·  Health monitoring',
         2.8),
        (AMBER, '#fef3c7', 'Layer 3 — React Frontend Dashboard',
         'Fetches data from API  ·  Renders 7 interactive charts\n'
         'Collects instructor interactions as study log events',
         1.1),
    ]

    for border, fill, title, body, y in layers:
        rect = mpatches.FancyBboxPatch((0.4, y), 9.2, 1.3,
            boxstyle='round,pad=0.08', linewidth=2,
            edgecolor=border, facecolor=fill)
        ax.add_patch(rect)
        ax.text(0.75, y + 1.0, title, fontsize=10, fontweight='bold', color=border, va='top')
        ax.text(0.75, y + 0.58, body,  fontsize=8.5, color=GREY, va='top', linespacing=1.5)

    # Arrows between layers
    for y_start, y_end, label in [(4.5, 4.1, 'JSON files'), (2.8, 2.4, 'HTTP / REST')]:
        ax.annotate('', xy=(5, y_end), xytext=(5, y_start),
            arrowprops=dict(arrowstyle='->', color=GREY, lw=1.8))
        ax.text(5.2, (y_start + y_end) / 2, label, fontsize=8, color=GREY, va='center')

    ax.set_title('ConceptGrade — Three-Tier Architecture', fontsize=12, fontweight='bold',
                 color=GREY, pad=10)
    return save_fig(fig, '_diag_architecture.png')


def diagram_kg_example() -> str:
    """
    Knowledge Graph ego-graph showing concept relationships around "Backpropagation".

    Design:
    - Central node (Backpropagation) at (5, 3.5) with larger font & border
    - Supporting nodes positioned at corners and edges to maximize spacing
    - Edges drawn to node perimeters (using numpy to avoid clipping)
    - Edge labels placed at midpoints with white background for contrast
    - Legend positioned at bottom-center (2 columns) to avoid node overlaps
    """
    fig, ax = plt.subplots(figsize=(10, 7))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')

    # Node positions: (x, y, border_color, fill_color)
    # Central node at (5, 3.5); supporting nodes pushed to corners and edges
    # to maximize visual separation and avoid label overlaps
    nodes = {
        'Backpropagation': (5.0, 3.5, '#1e40af', '#dbeafe'),  # centre (blue)
        'Gradient\nDescent': (1.8, 5.5, '#15803d', '#dcfce7'),  # top-left
        'Learning\nRate':    (1.8, 1.5, '#15803d', '#dcfce7'),  # bottom-left
        'Neural\nNetwork':   (8.2, 5.5, '#15803d', '#dcfce7'),  # top-right
        'Weight\nUpdate':    (8.2, 1.5, '#b45309', '#fef3c7'),  # bottom-right
        'Chain\nRule':       (5.0, 1.0, '#6b21a8', '#f3e8ff'),  # bottom-centre
    }

    # (src, dst, label, text_offset_x, text_offset_y)
    edges = [
        ('Gradient\nDescent', 'Backpropagation', 'PREREQUISITE FOR', -0.6,  0.25),
        ('Learning\nRate',    'Backpropagation', 'OPERATES ON',      -0.5, -0.25),
        ('Chain\nRule',       'Backpropagation', 'PREREQUISITE FOR',  0.6,  0.25),
        ('Backpropagation',   'Neural\nNetwork', 'IMPLEMENTS',        0.5,  0.25),
        ('Backpropagation',   'Weight\nUpdate',  'PRODUCES',          0.5, -0.25),
    ]
    edge_colors = {
        'PREREQUISITE FOR': '#7c3aed',
        'OPERATES ON':      '#d97706',
        'IMPLEMENTS':       '#059669',
        'PRODUCES':         '#2563eb',
    }

    # Node radii (central node larger for emphasis)
    R = {'Backpropagation': 0.85}

    def node_r(label):
        """Return the radius for a given node label (0.70 or 0.85)."""
        return R.get(label, 0.70)

    # Draw edges with proper perimeter anchoring
    # Arrow endpoints are calculated by moving inward from each node's centre
    # by its radius (unit vector × radius). This avoids cutting through node
    # boundaries and produces clean arrow-to-perimeter endpoints.
    import numpy as np
    for src, dst, rel, tx, ty in edges:
        sx, sy = nodes[src][:2]
        dx, dy = nodes[dst][:2]
        # Compute unit vector from source to destination
        vec = np.array([dx - sx, dy - sy])
        length = np.linalg.norm(vec)
        unit = vec / length
        # Shorten arrow endpoints by node radii to land on perimeters
        src_r = node_r(src)
        dst_r = node_r(dst)
        start = np.array([sx, sy]) + unit * src_r
        end   = np.array([dx, dy]) - unit * dst_r
        color = edge_colors.get(rel, GREY)
        # Draw arrow with consistent styling
        ax.annotate('', xy=(end[0], end[1]), xytext=(start[0], start[1]),
            arrowprops=dict(arrowstyle='->', color=color, lw=1.8,
                            mutation_scale=16))
        # Place label at edge midpoint with manual offset (tx, ty)
        # White background ensures label is readable over crossing edges
        mx, my = (sx + dx) / 2 + tx, (sy + dy) / 2 + ty
        ax.text(mx, my, rel, fontsize=8, color=color,
                ha='center', va='center', fontweight='bold',
                bbox=dict(fc='white', ec='none', alpha=0.75, pad=1.5))

    # Draw nodes on top of edges (zorder=3) with text labels above (zorder=4)
    # Central node gets larger border and bold text for visual emphasis
    for label, (x, y, border, fill) in nodes.items():
        r = node_r(label)
        is_central = label == 'Backpropagation'
        # Circle with solid fill and coloured border
        ax.add_patch(plt.Circle((x, y), r, color=fill, ec=border,
                                lw=3.0 if is_central else 1.8, zorder=3))
        # Centred text label, bold and larger for central node
        ax.text(x, y, label, ha='center', va='center',
                fontsize=9.5 if is_central else 8.5,
                fontweight='bold' if is_central else 'normal',
                color=border, zorder=4, linespacing=1.35)

    # Legend positioned at bottom-center (2 columns) to avoid overlapping nodes
    # Explains the colour coding used for node categories
    legend_items = [
        mpatches.Patch(fc='#dbeafe', ec='#1e40af', label='Central concept (selected)'),
        mpatches.Patch(fc='#dcfce7', ec='#15803d', label='Expected in rubric'),
        mpatches.Patch(fc='#fef3c7', ec='#b45309', label='Related concept'),
        mpatches.Patch(fc='#f3e8ff', ec='#6b21a8', label='Supporting concept'),
    ]
    ax.legend(handles=legend_items, loc='lower center', ncol=2, fontsize=9,
              framealpha=0.95, edgecolor='#d1d5db')

    ax.set_title('Example: Knowledge Graph Ego-Graph for "Backpropagation"',
                 fontsize=12, fontweight='bold', color=GREY, pad=12)
    return save_fig(fig, '_diag_kg_example.png')


def diagram_brushing_flow() -> str:
    """
    Two independent interaction flows with coloured band backgrounds.

    Design principle: In ConceptGrade's dashboard, instructor interactions with
    different charts (heatmap, radar) can simultaneously filter the Answer Panel,
    but these are independent flows—not a single sequential chain. Each flow
    has its own visual band (blue for heatmap, purple for radar) with:
    - Separate starting point (heatmap click vs. radar quartile select)
    - Independent action sequence (boxes in each flow)
    - Clear label stating independence
    - Gap between bands for visual separation

    Layout ensures no label-box overlaps:
    - Tall bands (3.2 units) provide breathing room
    - Labels positioned 0.3 units below band top, 0.7 units above box tops
    - Divider note placed in the 0.7-unit gap between flows
    """
    fig, ax = plt.subplots(figsize=(13, 8.0))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 8.0)
    ax.axis('off')

    BW, BH = 1.9, 1.7  # box width and height (consistent across flows)
    # Large arrows (lw=3.5, mutation_scale=28) remain visible at document scale
    ARROW = dict(arrowstyle='->', color='#374151', lw=3.5, mutation_scale=28)

    # Coordinate layout for two-band design:
    #   Flow 1 (blue band, y=4.4–7.6): 4 boxes at y=5.7 connected by arrows
    #   gap with divider text at y≈4.0
    #   Flow 2 (purple band, y=0.5–3.7): 2 boxes at y=1.9 connected by arrow

    FLOW1_Y   = 5.7
    FLOW1_TOP = 7.6
    FLOW1_BOT = 4.4

    FLOW2_Y   = 1.9
    FLOW2_TOP = 3.7
    FLOW2_BOT = 0.5

    # ── Flow 1 background band (blue) ────────────────────────────────
    # Light blue rounded rectangle (zorder=1) behind Flow 1 boxes
    ax.add_patch(mpatches.FancyBboxPatch(
        (0.15, FLOW1_BOT), 12.7, FLOW1_TOP - FLOW1_BOT,
        boxstyle='round,pad=0.05', lw=1.5, ec='#93c5fd', fc='#eff6ff', zorder=1))
    # Flow 1 label at top of band: positioned 0.3 units below band top (y=7.3)
    # This places it 0.7+ units above the highest box edge (y=6.55), ensuring clear separation
    ax.text(0.50, FLOW1_TOP - 0.30,
            'Flow 1 — Main Drill-Down',
            ha='left', va='center', fontsize=10.5, fontweight='bold', color='#1e40af')

    # ── Flow 2 background band (purple) ──────────────────────────────
    # Light purple rounded rectangle (zorder=1) behind Flow 2 boxes
    ax.add_patch(mpatches.FancyBboxPatch(
        (0.15, FLOW2_BOT), 12.7, FLOW2_TOP - FLOW2_BOT,
        boxstyle='round,pad=0.05', lw=1.5, ec='#c4b5fd', fc='#f5f3ff', zorder=1))
    # Flow 2 label at top of band: positioned 0.3 units below band top (y=3.4)
    # This places it 0.65+ units above the highest box edge (y=2.75), ensuring clear separation
    ax.text(0.50, FLOW2_TOP - 0.30,
            'Flow 2 — Radar Quartile Filter  (independent of Flow 1)',
            ha='left', va='center', fontsize=10.5, fontweight='bold', color='#6b21a8')

    # ── Flow 1 boxes: 4-step interaction sequence ─────────────────────
    # Spacing: 1.4 → 4.4 → 7.4 → 10.4 (gaps of 3.0, box width 1.9 → net gap 1.1)
    FLOW1_CX = [1.4, 4.4, 7.4, 10.4]
    flow1 = [
        ('#1e40af', '#dbeafe', 'Click Heatmap Cell',    'Selects concept + severity'),
        ('#0f766e', '#ccfbf1', 'Answer Panel Filters',   'Shows matching students'),
        ('#b45309', '#fef3c7', 'Click Student',          'Fetches XAI data'),
        ('#15803d', '#dcfce7', 'KG Panel Updates',       'Nodes turn green / red / grey'),
    ]

    # ── Flow 2 boxes: 2-step independent filter ──────────────────────
    # Spacing: 1.4 → 4.4 (same x-positions as Flow 1 for visual alignment)
    FLOW2_CX = [1.4, 4.4]
    flow2 = [
        ('#6b21a8', '#f3e8ff', 'Click Radar Quartile',   'Select score range Q1\u2013Q4'),
        ('#0f766e', '#ccfbf1', 'Answer Panel Re-filters', 'List updates to that quartile'),
    ]

    def draw_box(cx, cy, border, fill, title, sub):
        """
        Draw a single action box with rounded borders, coloured background, and text.

        Args:
            cx, cy: centre coordinates
            border: border colour (also used for title text)
            fill: background fill colour
            title: main action text (bold)
            sub: supporting description (smaller, grey)
        """
        # Rounded box (zorder=2) behind text
        ax.add_patch(mpatches.FancyBboxPatch(
            (cx - BW / 2, cy - BH / 2), BW, BH,
            boxstyle='round,pad=0.12', lw=2.5, ec=border, fc=fill, zorder=2))
        # Bold title text slightly above centre
        ax.text(cx, cy + 0.28, title, ha='center', va='center', fontsize=9.5,
                fontweight='bold', color=border, linespacing=1.3, zorder=3)
        # Smaller description text slightly below centre
        ax.text(cx, cy - 0.36, sub, ha='center', va='center', fontsize=8.5,
                color=GREY, linespacing=1.3, zorder=3)

    # Draw all boxes for both flows
    for cx, (b, f, t, s) in zip(FLOW1_CX, flow1):
        draw_box(cx, FLOW1_Y, b, f, t, s)
    for cx, (b, f, t, s) in zip(FLOW2_CX, flow2):
        draw_box(cx, FLOW2_Y, b, f, t, s)

    # Flow 1 arrows: connect 4 boxes in sequence (3 arrows total)
    # Large arrows (lw=3.5, mutation_scale=28) ensure visibility at document scale
    for i in range(3):
        ax.annotate('', xy=(FLOW1_CX[i + 1] - BW / 2, FLOW1_Y),
                    xytext=(FLOW1_CX[i] + BW / 2, FLOW1_Y),
                    arrowprops=ARROW, zorder=4)

    # Flow 2 arrow: connect 2 boxes in sequence (1 arrow total)
    ax.annotate('', xy=(FLOW2_CX[1] - BW / 2, FLOW2_Y),
                xytext=(FLOW2_CX[0] + BW / 2, FLOW2_Y),
                arrowprops=ARROW, zorder=4)

    # Divider note in the gap between flows (y ≈ 4.0)
    # Clarifies that the two flows are independent and both update the Answer Panel
    # This is the key message: there is NO link from box Flow1-4 to box Flow2-1
    ax.text(6.5, (FLOW1_BOT + FLOW2_TOP) / 2,
            '↕  These two flows are independent — both filter the Answer Panel separately.',
            ha='center', va='center', fontsize=9, color='#6b7280', style='italic')

    ax.set_title('Bidirectional Brushing \u2014 How Clicking One Chart Updates the Others',
                 fontsize=13, fontweight='bold', color=GREY, pad=14)
    return save_fig(fig, '_diag_brushing.png')


def diagram_pipeline() -> str:
    """
    Five-stage grading pipeline with wide box spacing.

    Design principle: The pipeline is a sequential 5-stage process where each
    stage takes the previous stage's output and produces refined grades.
    Spacing is generous (gap = 1.0 unit between boxes) so arrows remain
    clearly visible even when the figure is scaled down to fit in a Word document.
    """
    fig, ax = plt.subplots(figsize=(13, 4.5))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 4.5)
    ax.axis('off')

    # Layout parameters
    BW, BH = 1.6, 2.4  # box dimensions (width 1.6, height 2.4)
    BOT = 0.8          # baseline (y-coord of box bottoms)
    MID = BOT + BH / 2  # midline for arrow placement (y ≈ 1.9)
    # x-centres: spacing 2.6 units apart → gap = 2.6 - 1.6 = 1.0 per arrow
    CX  = [1.0, 3.6, 6.2, 8.8, 11.4]

    stages = [
        ('#1e40af', '#dbeafe', 'Stage 1', 'Self-Consistent\nExtractor'),
        ('#0f766e', '#ccfbf1', 'Stage 2', 'Confidence-\nWeighted\nComparator'),
        ('#b45309', '#fef3c7', 'Stage 3', 'LLM-as-\nVerifier'),
        ('#15803d', '#dcfce7', 'Stage 4', 'Chain Coverage\nScorer'),
        ('#7c3aed', '#f3e8ff', 'Stage 5', 'Final Score\n+ Explanation'),
    ]

    # Draw all 5 stage boxes
    for cx, (border, fill, label, name) in zip(CX, stages):
        # Rounded box (zorder=2) behind text
        ax.add_patch(mpatches.FancyBboxPatch(
            (cx - BW / 2, BOT), BW, BH,
            boxstyle='round,pad=0.12', lw=2.5, ec=border, fc=fill, zorder=2))
        # Stage label (e.g. "Stage 1") at top in bold
        ax.text(cx, BOT + BH - 0.28, label,
                ha='center', va='center', fontsize=11, fontweight='bold',
                color=border, zorder=3)
        # Stage name (e.g. "Self-Consistent Extractor") below label
        ax.text(cx, MID - 0.12, name,
                ha='center', va='center', fontsize=10, color=GREY,
                linespacing=1.5, zorder=3)

    # Connect all 5 boxes with 4 sequential arrows
    # Large arrows (lw=3.5, mutation_scale=28) remain visible at document scale
    ARROW = dict(arrowstyle='->', color='#374151', lw=3.5, mutation_scale=28)
    for i in range(4):
        # Arrow from right edge of stage i to left edge of stage i+1
        ax.annotate('', xy=(CX[i + 1] - BW / 2, MID),
                    xytext=(CX[i] + BW / 2, MID),
                    arrowprops=ARROW, zorder=4)

    ax.set_title('The Five-Stage ConceptGrade Grading Pipeline',
                 fontsize=13, fontweight='bold', color=GREY, pad=12)
    return save_fig(fig, '_diag_pipeline.png')


def diagram_study_conditions() -> str:
    """Side-by-side comparison of Condition A vs Condition B."""
    fig, (ax_a, ax_b) = plt.subplots(1, 2, figsize=(10, 4.5))
    fig.patch.set_facecolor(BG)

    def draw_condition(ax, title, color, items):
        ax.set_facecolor(WHITE)
        ax.set_xlim(0, 6)
        ax.set_ylim(0, len(items) * 0.9 + 0.8)
        ax.axis('off')
        ax.set_title(title, fontsize=11, fontweight='bold', color=color, pad=8)
        for i, (label, active) in enumerate(reversed(items)):
            y = i * 0.9 + 0.3
            fc = color if active else '#e5e7eb'
            ec = color if active else '#9ca3af'
            rect = mpatches.FancyBboxPatch((0.2, y), 5.6, 0.7,
                boxstyle='round,pad=0.05', lw=1.5, ec=ec, fc=fc + '33' if active else fc)
            ax.add_patch(rect)
            tc = color if active else '#9ca3af'
            ax.text(3.0, y + 0.35, label, ha='center', va='center',
                    fontsize=8.5, color=tc, fontweight='bold' if active else 'normal')

    items_a = [
        ('Summary Metric Cards (MAE, N, p-value)', True),
        ('Misconception Heatmap', False),
        ('Student Answer Panel', False),
        ('Knowledge Graph Panel', False),
        ('Score Samples + XAI Text', False),
    ]
    items_b = [
        ('Summary Metric Cards (MAE, N, p-value)', True),
        ('Misconception Heatmap', True),
        ('Student Answer Panel', True),
        ('Knowledge Graph Panel', True),
        ('Score Samples + XAI Text', True),
    ]

    draw_condition(ax_a, 'Condition A — Control\n(Numbers only)', RED, items_a)
    draw_condition(ax_b, 'Condition B — Treatment\n(Full Dashboard)', GREEN, items_b)

    fig.suptitle('User Study Conditions', fontsize=12, fontweight='bold', color=GREY, y=1.01)
    plt.tight_layout()
    return save_fig(fig, '_diag_conditions.png')


# ═══════════════════════════════════════════════════════════════════════════════
# WORD DOCUMENT HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def set_heading(doc: Document, text: str, level: int):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return h


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc: Document, text: str, level: int = 0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    p = doc.add_paragraph(text, style=style)
    return p


def add_image(doc: Document, path: str, width_inches: float = 5.8, caption: str = ''):
    """
    Add a centred image with formatted caption to the document.

    Caption format: "Figure N — Description text"
    - "Figure N" is rendered in bold
    - "— Description text" is rendered in italic
    - Spacing: 10pt before image, 14pt after caption for visual separation

    Args:
        doc: python-docx Document object
        path: file path to PNG image
        width_inches: image width (default 5.8", leaves ~1" margins on A4)
        caption: text in format "Figure N — Description" or "Figure N - Description"
    """
    # Breathing room before the figure block (10pt space)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Inches(width_inches))

    if caption:
        # Caption paragraph (centred, smaller font)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(14)

        # Parse caption: split on em-dash or hyphen to separate label from description
        # "Figure N — Description" → label="Figure N" (bold), title="Description" (italic)
        if ' \u2014 ' in caption:
            label, _, title = caption.partition(' \u2014 ')
        elif ' - ' in caption:
            label, _, title = caption.partition(' - ')
        else:
            label, title = caption, ''

        # Label in bold (e.g. "Figure 1")
        bold_run = cap.add_run(label)
        bold_run.font.bold = True
        bold_run.font.size = Pt(9)
        bold_run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

        if title:
            sep_run = cap.add_run(' \u2014 ')
            sep_run.font.size = Pt(9)
            sep_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

            desc_run = cap.add_run(title)
            desc_run.font.italic = True
            desc_run.font.size = Pt(9)
            desc_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)


def add_toc(doc: Document):
    """Insert a Word TOC field. Word/LibreOffice must update fields on open (Ctrl+A, F9)."""
    para = doc.add_paragraph()
    run = para.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar_begin)
    run._r.append(instrText)
    run._r.append(fldChar_separate)
    run._r.append(fldChar_end)

    placeholder = doc.add_paragraph('[ Right-click here → Update Field to generate the Table of Contents ]')
    placeholder.runs[0].font.italic = True
    placeholder.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table(doc: Document, headers: list, rows: list,
              col_widths: list = None, header_color=None):
    """Add a styled table with bold headers and alternating row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Header background
        tc_pr = hdr_cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        color = (header_color or '1e40af').lstrip('#')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tc_pr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        fill = 'f0f9ff' if r_idx % 2 == 0 else 'ffffff'
        for c_idx, cell_text in enumerate(row_data):
            cells[c_idx].text = str(cell_text)
            cells[c_idx].paragraphs[0].runs[0].font.size = Pt(9)
            tc_pr = cells[c_idx]._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tc_pr.append(shd)

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_info_box(doc: Document, text: str, color_hex: str = '1e40af'):
    """Add a tinted paragraph that acts as a callout box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.italic = True
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16))
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '4')
    left.set(qn('w:color'), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

def build_document():
    """
    Generate and assemble the complete ConceptGrade System Overview Word document.

    Steps:
    1. Generate all 5 matplotlib diagrams as PNG files in the current directory
    2. Create a Word document with proper style hierarchy
    3. Add table of contents (auto-generated from heading styles)
    4. Populate content sections: overview, problem statement, architecture, design details
    5. Embed diagrams with proper captions and spacing
    6. Add evaluation tables and results summaries
    7. Save as ConceptGrade_System_Overview.docx
    """
    print("Generating diagrams...")
    # Generate all diagrams as PNG files (saved to OUT_DIR)
    img_arch       = diagram_architecture()
    img_kg         = diagram_kg_example()
    img_brushing   = diagram_brushing_flow()
    img_pipeline   = diagram_pipeline()
    img_conditions = diagram_study_conditions()
    print("Diagrams saved.")

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Default style tweaks ──────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for lvl in range(1, 4):
        h = doc.styles[f'Heading {lvl}']
        h.font.name = 'Calibri'
        colors = ['1e40af', '0f766e', '374151']
        r, g, b = [int(colors[lvl-1][i:i+2], 16) for i in (0, 2, 4)]
        h.font.color.rgb = RGBColor(r, g, b)
        h.font.size = Pt([18, 14, 12][lvl - 1])
        h.paragraph_format.space_before = Pt([18, 12, 8][lvl - 1])
        h.paragraph_format.space_after  = Pt([8,  6,  4][lvl - 1])

    # ══════════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('ConceptGrade')
    title_run.font.size = Pt(32)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(
        'A Visual Analytics Dashboard for\nKnowledge Graph-Grounded Automated Essay Grading')
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x0f, 0x76, 0x6e)

    doc.add_paragraph()

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run('System Overview and Implementation Documentation')
    meta_run.font.size = Pt(12)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.add_run('Brahmaji Katragadda').font.size = Pt(12)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.add_run('April 2026').font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, 'Table of Contents', 1)
    add_toc(doc)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — WHAT PROBLEM DOES THIS SOLVE?
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '1. What Problem Does This Solve?', 1)

    add_body(doc,
        'When a student submits a written answer to an exam or homework question, grading it '
        'fairly and quickly is hard. In large classes, instructors can have hundreds of answers '
        'to read. Automated grading tools exist, but most of them just give a number — a score '
        '— without explaining why that score was given. This makes it difficult for instructors '
        'to trust the system, and even harder to understand which specific concepts a student is missing.')

    add_body(doc, 'ConceptGrade addresses this in two ways:')

    add_bullet(doc,
        'It grades student answers using a Knowledge Graph (KG) — a structured map of the '
        'important concepts in a subject and how they relate to each other. Instead of just '
        'comparing words, the system checks whether a student\'s answer demonstrates '
        'understanding of the right concepts and their relationships.')
    add_bullet(doc,
        'It shows instructors an interactive visual dashboard where they can explore grading '
        'results, see which concepts an entire class is struggling with, drill down into '
        'individual student answers, and verify whether the automated score makes sense.')

    add_info_box(doc,
        'The goal is not to replace the instructor. It is to give them a tool that makes it '
        'faster and easier to identify where students are going wrong — and to check the '
        'automated scores critically rather than just trusting them blindly.',
        '0f766e')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — SYSTEM ARCHITECTURE
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '2. System Architecture', 1)

    add_body(doc,
        'The system is built in three layers that work together. Each layer has a clearly '
        'defined responsibility, and they communicate through well-defined interfaces.')

    add_image(doc, img_arch, width_inches=5.8,
              caption='Figure 1 — ConceptGrade Three-Tier Architecture')

    set_heading(doc, 'Why Three Layers?', 2)
    add_bullet(doc,
        'The Python pipeline is where the machine learning and Knowledge Graph work happens. '
        'It runs offline (not in real time) and saves its results to JSON files.')
    add_bullet(doc,
        'The NestJS backend is a lightweight bridge that serves those files to the browser. '
        'It does not run machine learning — it reads and organises pre-computed data.')
    add_bullet(doc,
        'The React frontend is what the instructor sees. It fetches data from the backend '
        'and presents it as interactive charts.')

    add_body(doc,
        'This separation of concerns means the machine learning code and the visual interface '
        'are completely independent. Changes to the grading algorithm do not affect the '
        'dashboard, and vice versa.')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — KNOWLEDGE GRAPH AND GRADING PIPELINE
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '3. The Knowledge Graph and Grading Pipeline', 1)

    set_heading(doc, '3.1  What Is a Knowledge Graph?', 2)

    add_body(doc,
        'For each exam question, the system builds a small graph where nodes are important '
        'concepts and edges are the relationships between them. The graph captures not just '
        'what concepts exist, but how they depend on each other. This is important for '
        'understanding gaps — if a student does not understand gradient descent, they probably '
        'cannot properly explain backpropagation either.')

    add_image(doc, img_kg, width_inches=5.2,
              caption='Figure 2 — Example Knowledge Graph Ego-Graph for the concept "Backpropagation"')

    add_body(doc, 'The system supports eight types of relationships between concepts:')
    add_table(doc,
        headers=['Relationship Type', 'Meaning'],
        rows=[
            ['PREREQUISITE_FOR', 'Concept A must be understood before Concept B'],
            ['PRODUCES',         'Concept A leads to or creates Concept B'],
            ['HAS_PART',         'Concept B is a component of Concept A'],
            ['HAS_PROPERTY',     'Concept B is a property or attribute of Concept A'],
            ['IMPLEMENTS',       'Concept A is a practical realisation of Concept B'],
            ['OPERATES_ON',      'Concept A acts upon or uses Concept B'],
            ['CONTRASTS_WITH',   'Concept A and Concept B are opposites or alternatives'],
            ['VARIANT_OF',       'Concept A is a variation or extension of Concept B'],
        ],
        col_widths=[2.2, 4.0],
        header_color='1e40af')

    set_heading(doc, '3.2  The Five-Stage Grading Pipeline', 2)

    add_body(doc,
        'The grading system runs five processing steps on each student answer. Each stage '
        'adds more information until a final score and explanation are produced.')

    add_image(doc, img_pipeline, width_inches=6.2,
              caption='Figure 3 — The Five-Stage ConceptGrade Grading Pipeline')

    add_table(doc,
        headers=['Stage', 'Name', 'What It Does'],
        rows=[
            ['1', 'Self-Consistent Extractor',
             'Extracts concepts from the answer three times independently. Only concepts that appear consistently across all three attempts are kept.'],
            ['2', 'Confidence-Weighted Comparator',
             'Compares the extracted concepts against the expected concepts listed in the Knowledge Graph. Each match is weighted by confidence.'],
            ['3', 'LLM-as-Verifier',
             'Uses three different AI "personas" to cross-check concept matches — like having three independent reviewers vote on each concept.'],
            ['4', 'Chain Coverage Scorer',
             'Calculates how much of the prerequisite chain of concepts the student demonstrated, not just isolated concepts.'],
            ['5', 'Final Score + Explanation',
             'Combines all signals into a final score with a plain-language explanation of which concepts were matched and which were missed.'],
        ],
        col_widths=[0.6, 2.0, 3.6],
        header_color='0f766e')

    set_heading(doc, '3.3  Results Across Three Datasets', 2)

    add_body(doc,
        'The pipeline was tested against a baseline (a plain language model without the '
        'Knowledge Graph). Lower MAE (Mean Absolute Error) means the score is closer to what '
        'a human expert would give:')

    add_table(doc,
        headers=['Dataset', 'Subject', 'Students', 'Baseline Error', 'Our System', 'Improvement'],
        rows=[
            ['Mohler 2011',   'Computer Science',    '120',  '0.330', '0.223', '▼ 32.4%  (p = 0.001)'],
            ['DigiKlausur',   'Neural Networks',     '646',  '0.394', '0.296', '▼ 24.9%  (p = 0.049)'],
            ['Kaggle ASAG',   'Elementary Science',  '473',  '0.244', '0.252', '→ Minimal change (n.s.)'],
        ],
        col_widths=[1.4, 1.6, 0.9, 1.1, 1.1, 1.8],
        header_color='15803d')

    add_info_box(doc,
        'The improvement is strongest in technical domains (Computer Science, Neural Networks) '
        'where subject-specific vocabulary is important. In general-language domains '
        '(elementary science), the Knowledge Graph provides less advantage because everyday '
        'vocabulary is not as structured.',
        '15803d')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — THE DASHBOARD
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '4. The Dashboard — What Was Built', 1)

    add_body(doc,
        'The dashboard has seven visual components. They are all connected: clicking something '
        'in one chart automatically updates the others. This kind of connection is called '
        'Coordinated Multiple Views (CMV). The design follows a well-known visual analytics '
        'principle: "overview first, zoom and filter, then details on demand."')

    add_table(doc,
        headers=['Component', 'Purpose', 'Key Interaction'],
        rows=[
            ['Misconception Heatmap',
             'Shows which concepts the whole class struggles with',
             'Click a cell → filters the Answer Panel'],
            ['Student Answer Panel',
             'Shows individual student answers with scores',
             'Click a student → updates the KG Panel'],
            ['Radar Chart',
             'Shows score profiles across cognitive dimensions',
             'Click a quartile chip → filters the Answer Panel'],
            ['KG Ego-Graph',
             'Shows the concept neighbourhood and student coverage',
             'Drag nodes · hover for details · colours update per student'],
            ['Score Samples Table',
             'Shows individual scores with XAI concept explanations',
             'Expand a row → see matched and missed concept chips'],
            ['Cross-Dataset Slopegraph',
             'Compares system improvement across all three datasets',
             'Read-only reference chart'],
            ['Class Summary Cards',
             'Headline metrics: N, MAE, improvement %, p-value',
             'Read-only reference cards'],
        ],
        col_widths=[1.8, 2.6, 2.4],
        header_color='1e40af')

    set_heading(doc, '4.1  Misconception Heatmap', 2)
    add_body(doc,
        'A grid where rows are concepts and columns are severity levels (Critical, Moderate, '
        'Minor, Covered). Each cell is coloured based on how many students fall into that '
        'category — darker means more students affected. Clicking a cell automatically filters '
        'the Student Answer Panel to show only the relevant students.')

    set_heading(doc, '4.2  Student Answer Panel (Master-Detail Layout)', 2)
    add_body(doc,
        'A two-column layout. The left column shows a scrollable list of students with score '
        'badges and colour-coded severity chips. The right column shows the full detail for '
        'the selected student — their complete answer, the reference answer, their SOLO '
        'level (how structurally complex their answer is), and their Bloom\'s level (what '
        'cognitive skill they demonstrated). Clicking a student in the list also triggers '
        'an asynchronous fetch of their XAI concept data, which updates the KG Panel.')

    set_heading(doc, '4.3  Radar Chart (Score Quartile Filter)', 2)
    add_body(doc,
        'A radar chart showing average performance across five cognitive dimensions — concept '
        'coverage, chain coverage, SOLO level, Bloom level, and score accuracy — broken into '
        'four quartiles (Q1 = bottom 25%, Q4 = top 25%). Clicking a quartile chip filters '
        'the Student Answer Panel to show only students in that score range.')

    set_heading(doc, '4.4  Knowledge Graph Ego-Graph', 2)
    add_body(doc,
        'An interactive diagram drawn as SVG (Scalable Vector Graphics) showing the selected '
        'concept at the centre with related concepts arranged around it. Edges are labelled '
        'with relationship types. Instructors can drag nodes to rearrange the layout, hover '
        'over nodes and edges for full descriptions, and see a student overlay where each '
        'node turns green (demonstrated), red (expected but missing), or grey (not relevant).')

    set_heading(doc, '4.5  Score Samples Table with XAI Explanation', 2)
    add_body(doc,
        'A table of individual student scores with an expand button on each row. When expanded, '
        'the row shows which concepts the student matched (green chips), which they missed '
        '(red chips), and a causal text explanation explaining why the score changed. This is '
        'the explainability layer — it tells the instructor not just the score, but why.')

    set_heading(doc, '4.6  Cross-Dataset Slopegraph', 2)
    add_body(doc,
        'A line chart comparing the system\'s error against the baseline across all three '
        'datasets. Two bracket annotations group the datasets by type — Academic Domains '
        '(Computer Science, Neural Networks) and General Education (Elementary Science) — '
        'with a subtitle observing that vocabulary richness correlates with improvement.')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — COORDINATED VIEWS AND BRUSHING
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '5. How the Charts Are Connected (Bidirectional Brushing)', 1)

    add_body(doc,
        'All chart components share a single piece of shared state called DashboardContext. '
        'This means that when one chart updates its selection, all other relevant charts '
        'update automatically. The diagram below shows the two main interaction flows:')

    add_image(doc, img_brushing, width_inches=6.2,
              caption='Figure 4 — Bidirectional Brushing: How Clicking One Chart Updates the Others')

    add_body(doc,
        'The top row shows the main drill-down flow an instructor follows to go from class-wide '
        'patterns to a specific student\'s conceptual gap. The bottom row shows how the Radar '
        'Chart\'s quartile filter applies independently, narrowing the student list to a '
        'specific score band while keeping the concept selection in place.')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — SHARED STATE
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '6. The Shared State System (DashboardContext)', 1)

    add_body(doc,
        'All interactive variables are stored in one central place — called DashboardContext '
        '— that every component can read from and write to simultaneously:')

    add_table(doc,
        headers=['Variable', 'What It Tracks'],
        rows=[
            ['selectedConcept',              'Which concept was clicked in the heatmap'],
            ['selectedSeverity',             'Which severity column was clicked (Critical / Moderate / Minor)'],
            ['selectedStudentId',            'Which student is currently selected in the Answer Panel'],
            ['selectedStudentMatchedConcepts', 'Which concept IDs that student covered (from XAI fetch)'],
            ['studentOverlayLoading',        'Whether the KG node colours are currently being fetched'],
            ['studentOverlayError',          'Whether the most recent KG overlay fetch failed'],
            ['selectedQuartileIndex',        'Which radar quartile filter is currently active (Q1–Q4)'],
        ],
        col_widths=[2.8, 3.9],
        header_color='7c3aed')

    add_body(doc,
        'The state is managed using React\'s useReducer pattern — all state transitions are '
        'defined as named actions (SELECT_CONCEPT, SELECT_STUDENT, SET_LOADING, SET_ERROR, '
        'SELECT_QUARTILE, CLEAR_ALL). This makes the logic predictable: every place in the '
        'code that changes state does so by dispatching a named action, not by directly '
        'modifying variables.')

    add_info_box(doc,
        'Why useReducer instead of useState? Several actions have side effects on other fields '
        '— for example, selecting a new concept must also clear the previously selected '
        'student. A reducer puts all of this clearing logic in one place, making it impossible '
        'to forget a step.',
        '7c3aed')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — TECHNICAL DECISIONS
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '7. Key Technical Decisions and Why They Were Made', 1)

    decisions = [
        ('7.1  Async File Reading with a Cache',
         'The backend originally used synchronous file reading, which blocks the server from '
         'handling any other requests while a file is being read. This is fine for a single '
         'user but causes noticeable delays when multiple users access the system simultaneously.',
         'Replaced all file reads with asynchronous reads (fs.promises.readFile) and added '
         'an in-memory cache. The first time a dataset is requested the file is read and '
         'stored in memory. Every subsequent request for the same file is served from memory '
         'without touching the disk.'),
        ('7.2  Race Condition Guard in the KG Overlay Fetch',
         'When an instructor clicks quickly through several students, multiple network requests '
         'go out at the same time. If request 3 comes back before request 2, the wrong '
         'student\'s concepts would appear in the KG panel.',
         'A reference (latestSelectedIdRef) tracks which student was most recently clicked. '
         'When a fetch result comes back, it is only applied if the student ID still matches '
         'the currently selected student. Stale results from earlier clicks are discarded.'),
        ('7.3  KG Drag Cursor',
         'While dragging a node, the cursor should change to a "grabbing" hand icon. The '
         'initial implementation stored the drag state in a React ref, but reading a ref in '
         'JSX does not trigger a re-render, so the cursor never visually updated.',
         'The cursor style is set directly on the SVG element\'s DOM property '
         '(svgRef.current.style.cursor) on mousedown and reset on mouseup. This bypasses '
         'React\'s rendering cycle for this high-frequency visual update.'),
        ('7.4  Path Sanitisation',
         'The API endpoints accept a :dataset parameter in the URL that is used to construct '
         'a file path on the server. Without validation, a crafted URL could attempt to read '
         'files outside the intended data directory (path traversal attack).',
         'path.basename(dataset) is applied to every incoming dataset parameter before it '
         'is used in a file path. This strips any directory components, ensuring only a '
         'filename — not a full path — reaches the file system.'),
        ('7.5  Knowledge Graph Edge Weight Averaging',
         'The same two concepts might be related in multiple questions, each time with a '
         'slightly different relationship strength. The original code kept only the first '
         'occurrence and discarded all others.',
         'When the same edge appears multiple times, the weights are accumulated and '
         'averaged. The final edge weight represents the consensus strength across all '
         'questions where that relationship appears.'),
        ('7.6  Question-Scoped Expected Concepts',
         'Initially, a node in the KG panel was marked as "expected" (red = missing) if it '
         'was required in any question across the dataset. This produced false signals — '
         'a concept would appear red even if it was never required for the question the '
         'student actually answered.',
         'The KG endpoint now accepts an optional ?questionId= query parameter. When '
         'provided, a node is only marked as expected if it is required for that specific '
         'question. This makes the red/green visual encoding pedagogically accurate.'),
        ('7.7  CORS Configuration',
         'During the user study, the React frontend (port 5173) needs to talk to the NestJS '
         'backend (port 5001). Browsers block this by default unless the server explicitly '
         'permits it (Cross-Origin Resource Sharing — CORS).',
         'Added http://localhost:5173 and http://127.0.0.1:5173 to the backend\'s list of '
         'allowed origins, so the browser permits fetch requests during a local study session.'),
    ]

    for title, problem, solution in decisions:
        set_heading(doc, title, 2)
        add_body(doc, '⚠  Problem: ' + problem)
        add_body(doc, '✓  Solution: ' + solution)

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — USER STUDY INFRASTRUCTURE
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '8. User Study Infrastructure', 1)

    add_body(doc,
        'The dashboard is designed not just as a tool but as a research instrument for '
        'measuring how instructors use it. Two study conditions are supported:')

    add_image(doc, img_conditions, width_inches=5.6,
              caption='Figure 5 — Condition A (Control) vs Condition B (Treatment)')

    set_heading(doc, '8.1  Study Conditions', 2)
    add_body(doc,
        'The condition is set via a URL parameter (?condition=A or ?condition=B). The same '
        'codebase serves both conditions, ensuring the only difference between participants '
        'is the presence or absence of the visual tools.')

    add_table(doc,
        headers=['Condition', 'What the Instructor Sees', 'Purpose'],
        rows=[
            ['A — Control',
             'Four summary metric cards only (N, MAE, improvement %, p-value). No charts.',
             'Represents the current state of practice: instructors receive automated scores with no visual tools.'],
            ['B — Treatment',
             'The full interactive dashboard with all seven chart components and linking.',
             'Tests whether the visual tools change how instructors explore and verify grading results.'],
        ],
        col_widths=[1.2, 2.8, 2.7],
        header_color='b45309')

    set_heading(doc, '8.2  Study Task Panel', 2)
    add_body(doc,
        'When study mode is active, a task panel appears at the top of the dashboard. It '
        'shows the instructor a diagnostic task prompt, a text box for their written response, '
        'a confidence slider (1–5), and a Submit button. The panel records the time between '
        'first focus and submission.')

    set_heading(doc, '8.3  Study Event Logging (Two Storage Layers)', 2)
    add_body(doc,
        'Every interaction is recorded as a timestamped event. Two storage layers ensure '
        'events are not lost if the browser closes unexpectedly:')

    add_table(doc,
        headers=['Event Type', 'When It Is Recorded'],
        rows=[
            ['page_view',   'When the dashboard first loads'],
            ['tab_change',  'When the instructor switches between datasets'],
            ['task_start',  'When the instructor first clicks into the response text box'],
            ['task_submit', 'When the instructor submits (includes answer, confidence, time taken)'],
            ['chart_hover', 'When the instructor interacts with a chart'],
        ],
        col_widths=[1.6, 5.1],
        header_color='0f766e')

    add_body(doc, 'Storage layers:')
    add_bullet(doc,
        'localStorage (primary) — events are written to the browser immediately. '
        'The instructor can export all events as a JSON file at any time.')
    add_bullet(doc,
        'Server-side JSONL backup (secondary) — each event is also sent to the backend '
        '(POST /api/study/log) and appended to a per-session file '
        '(data/study_logs/{session_id}.jsonl). Even if the instructor closes the browser '
        'without exporting, the data is safe on the server.')

    set_heading(doc, '8.4  Backend Health Check', 2)
    add_body(doc,
        'The backend exposes a GET /api/study/health endpoint. The frontend pings this on '
        'every page load. If the backend is unreachable, a red warning banner appears at '
        'the top of the dashboard, alerting the study facilitator before they begin a '
        'participant session. This prevents silent data loss.')

    set_heading(doc, '8.5  XAI Overlay Error Indicator', 2)
    add_body(doc,
        'When an instructor clicks a student, the system fetches the student\'s XAI concept '
        'data to colour the KG nodes. If this fetch fails, the KG panel shows a small amber '
        'chip in its header reading "⚠ overlay unavailable". This chip is persistent — it '
        'stays visible until the instructor clicks a different student. A persistent chip '
        'was chosen over a timed pop-up notification because a notification that disappears '
        'automatically can interrupt an instructor\'s verbal explanation during a think-aloud '
        'study session.')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 9 — ITERATIVE DESIGN
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '9. The Iterative Design Process', 1)

    add_body(doc,
        'The dashboard was built through four structured rounds of critique and improvement. '
        'Each round identified specific problems and implemented targeted fixes, following '
        'the Design Study Methodology (DSM) — a recognised approach in visual analytics '
        'research where design decisions are driven by domain expert feedback.')

    add_table(doc,
        headers=['Round', 'Focus', 'Key Changes Made'],
        rows=[
            ['v1 → v2', 'Core interactivity',
             'Added heatmap-to-panel linking, KG ego-graph, radar quartile filter, cross-dataset slopegraph'],
            ['v2 → v3', 'Unified linking',
             'Unified KG overlay so clicking a student from either the Answer Panel or Score Table updates the KG; added slopegraph domain annotations'],
            ['v3 → v4', 'Polish and robustness',
             'KG loading indicator with node dimming, print-safe chart colours, outcome-oriented onboarding tips'],
            ['v4 → v5', 'Study design',
             'Finalised three-task study protocol; identified key think-aloud moments to watch for'],
        ],
        col_widths=[0.8, 1.6, 4.3],
        header_color='7c3aed')

    add_body(doc,
        'After the dashboard was declared feature-complete, a code-level review identified '
        'and fixed additional technical issues:')

    add_table(doc,
        headers=['Issue Found', 'Fix Applied'],
        rows=[
            ['Synchronous file I/O blocking the server',           'Replaced with async reads + in-memory cache'],
            ['Drag cursor not updating during node drag',          'Direct DOM mutation via ref instead of React state'],
            ['localhost blocked by browser CORS policy',           'Added localhost to backend allowed origins'],
            ['Complex state logic scattered across components',    'Migrated to useReducer with named action types'],
            ['Path traversal vulnerability in API parameters',     'path.basename() applied to all URL parameters used in file paths'],
            ['KG nodes incorrectly marked as expected',            'Added question-scoped is_expected via ?questionId= parameter'],
            ['Edge weights lost on deduplication',                 'Changed to average weights across all occurrences'],
            ['Silent failure on XAI fetch error',                  'Added persistent ⚠ overlay unavailable chip in KG panel'],
            ['No server-side study log backup',                    'Added POST /api/study/log endpoint writing to JSONL files'],
            ['No health visibility for study facilitator',         'Added GET /api/study/health + frontend warning banner'],
        ],
        col_widths=[3.1, 3.6],
        header_color='374151')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 10 — FILE STRUCTURE
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '10. File Structure', 1)

    add_body(doc,
        'The repository is organised as a monorepo with three packages, each corresponding '
        'to one tier of the architecture:')

    add_table(doc,
        headers=['File / Folder', 'Purpose'],
        rows=[
            ['packages/concept-aware/',                   'Python pipeline, KG generation, evaluation results'],
            ['  data/*_eval_results.json',                'Pre-computed grading results for each dataset'],
            ['  data/*_auto_kg.json',                     'Auto-generated Knowledge Graphs per dataset'],
            ['  data/study_logs/*.jsonl',                 'Server-side study event logs (one file per session)'],
            ['packages/backend/src/visualization/',      'NestJS module: REST endpoints for datasets and KG'],
            ['  visualization.controller.ts',             'Route definitions (7 GET endpoints)'],
            ['  visualization.service.ts',               'File I/O, caching, data transformation logic'],
            ['  visualization.types.ts',                 'TypeScript interfaces shared between layers'],
            ['packages/backend/src/study/',              'NestJS module: study logging and health'],
            ['  study.controller.ts',                    'POST /api/study/log · GET /api/study/health'],
            ['  study.service.ts',                       'JSONL file writing logic'],
            ['packages/frontend/src/contexts/',          'React shared state'],
            ['  DashboardContext.tsx',                   'useReducer state machine for all CMV selections'],
            ['packages/frontend/src/components/charts/', '7 chart components'],
            ['  MisconceptionHeatmap.tsx',               'Heatmap grid component'],
            ['  StudentAnswerPanel.tsx',                 'Master-detail student list + detail pane'],
            ['  StudentRadarChart.tsx',                  'Quartile radar chart with filter chips'],
            ['  ConceptKGPanel.tsx',                     'SVG ego-graph with drag, hover, and student overlay'],
            ['  ScoreSamplesTable.tsx',                  'Expandable score table with XAI concept chips'],
            ['  CrossDatasetComparisonChart.tsx',        'SVG slopegraph with domain brackets'],
            ['packages/frontend/src/pages/',            'Top-level pages'],
            ['  InstructorDashboard.tsx',               'Main layout, condition gating, study task panel'],
            ['packages/frontend/src/utils/',            'Utilities'],
            ['  studyLogger.ts',                        'Event logging to localStorage + server backup'],
        ],
        col_widths=[3.1, 3.6],
        header_color='1e40af')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 11 — HOW TO RUN
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '11. How to Run the System', 1)

    add_body(doc, 'Prerequisites: Node.js 18+, Python 3.12, Yarn')

    set_heading(doc, 'Start the Backend', 2)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    r = p.add_run(
        'cd packages/backend\n'
        'node_modules/.bin/nest start\n'
        '# API available at http://localhost:5001')
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

    set_heading(doc, 'Start the Frontend', 2)
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    r = p.add_run(
        'cd packages/frontend\n'
        'yarn dev\n'
        '# Dashboard available at http://localhost:5173')
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

    set_heading(doc, 'Dashboard URLs', 2)
    add_table(doc,
        headers=['URL', 'What You See'],
        rows=[
            ['http://localhost:5173/dashboard',            'Full dashboard (default)'],
            ['http://localhost:5173/dashboard?condition=A','Control — metric cards only (no charts)'],
            ['http://localhost:5173/dashboard?condition=B','Treatment — full dashboard + study task panel'],
            ['http://localhost:5001/api/study/health',     'Backend health check (returns { status: "ok" })'],
        ],
        col_widths=[3.4, 3.3],
        header_color='374151')

    # ══════════════════════════════════════════════════════════════════════════
    # SECTION 12 — SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    set_heading(doc, '12. Summary', 1)

    add_body(doc,
        'ConceptGrade is a three-tier research system that combines automated essay grading '
        'with an interactive visual dashboard for instructors. The grading pipeline uses a '
        'Knowledge Graph to ground its scoring in structured domain knowledge, achieving '
        'significant error reduction in technical subject areas (32.4% in Computer Science, '
        '24.9% in Neural Networks).')

    add_body(doc,
        'The dashboard presents grading results through seven coordinated visual components '
        'that let instructors explore class-level patterns, drill down into individual students, '
        'verify automated scores using plain-language explanations, and build a richer '
        'understanding of where students are going wrong conceptually.')

    add_body(doc,
        'The system was developed through four rounds of structured design review, with each '
        'round driven by domain-expert feedback. A subsequent code-quality review addressed '
        'ten technical issues — covering performance, reliability, security, and study data '
        'integrity — resulting in a codebase that is robust and production-ready for a '
        'controlled user study.')

    add_info_box(doc,
        'The ultimate goal of ConceptGrade is not to automate the instructor out of the loop. '
        'It is to give instructors a sharper lens — so they can spend less time reading '
        'every answer and more time helping the students who need them most.',
        '1e40af')

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = os.path.join(OUT_DIR, 'ConceptGrade_System_Overview.docx')
    doc.save(out_path)
    print(f'\nDocument saved: {out_path}')

    # Clean up temporary diagram files
    for name in ['_diag_architecture.png', '_diag_kg_example.png',
                 '_diag_brushing.png', '_diag_pipeline.png', '_diag_conditions.png']:
        p = os.path.join(OUT_DIR, name)
        if os.path.exists(p):
            os.remove(p)
    print('Temporary diagram files removed.')


if __name__ == '__main__':
    build_document()
