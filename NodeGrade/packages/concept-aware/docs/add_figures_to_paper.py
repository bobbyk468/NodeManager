"""
Add figures to IEEE VIS paper - simplified approach using python-docx.
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load existing document
doc_path = 'IEEE_VIS_ConceptGrade_Complete_Paper.docx'
print(f"Loading {doc_path}...")
doc = Document(doc_path)

# Get current working directory
cwd = os.path.dirname(__file__)

# Figure insertion function
def insert_figure_after_section(doc, section_keyword, fig_path, fig_num, fig_caption):
    """Insert figure after a section heading that contains keyword."""
    for i, para in enumerate(doc.paragraphs):
        if section_keyword in para.text:
            # Find insertion point
            insert_idx = i + 2
            
            # Create new paragraph for image
            p_img = doc.paragraphs[insert_idx]._element
            new_p_img = doc.add_paragraph()
            new_p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            new_p_img.paragraph_format.space_before = Pt(10)
            new_p_img.paragraph_format.space_after = Pt(2)
            
            # Add image
            if os.path.exists(fig_path):
                run = new_p_img.add_run()
                run.add_picture(fig_path, width=Inches(5.8))
                print(f"  ✓ Inserted Figure {fig_num}")
            
            # Add caption
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(3)
            p_cap.paragraph_format.space_after = Pt(12)
            
            bold_run = p_cap.add_run(f'Figure {fig_num}')
            bold_run.font.bold = True
            bold_run.font.size = Pt(9)
            
            sep_run = p_cap.add_run(' — ')
            sep_run.font.size = Pt(9)
            
            desc_run = p_cap.add_run(fig_caption)
            desc_run.font.italic = True
            desc_run.font.size = Pt(9)
            
            return True
    return False

# Insert figures at strategic locations
print("Inserting figures...")

insert_figure_after_section(doc, '3. SYSTEM ARCHITECTURE',
    os.path.join(cwd, '_fig_architecture.png'), 1,
    'Three-tier ConceptGrade Architecture: Python Pipeline → NestJS API → React Dashboard')

insert_figure_after_section(doc, '3.2 Knowledge Graph Construction',
    os.path.join(cwd, '_fig_kg.png'), 2,
    'Knowledge Graph Example: Enzyme Kinetics Concepts and Relationships')

insert_figure_after_section(doc, '3.3 Five-Stage Grading Pipeline',
    os.path.join(cwd, '_fig_pipeline.png'), 3,
    'Five-Stage Grading Pipeline: From Student Answer to Score + Explanation')

insert_figure_after_section(doc, '4.2 Dashboard Components',
    os.path.join(cwd, '_fig_dashboard_layout.png'), 4,
    'Visual Analytics Dashboard Layout with 9 Linked Charts and Interactive Views')

insert_figure_after_section(doc, '4.3 Linked Views and Brushing',
    os.path.join(cwd, '_fig_brushing_flows.png'), 5,
    'Brushing Interaction Flows: Two Independent Workflows, Both Update Answer Panel')

insert_figure_after_section(doc, '6.1 Overall Performance',
    os.path.join(cwd, '_fig_results.png'), 6,
    'Results Comparison: MAE Reduction (%) and Statistical Significance (Wilcoxon Test)')

# Save new document
output_path = 'IEEE_VIS_ConceptGrade_Complete_WITH_FIGURES.docx'
doc.save(output_path)

print(f"\n✓ Paper with embedded figures saved: {output_path}")
print(f"  File size: {os.path.getsize(output_path) / 1024:.1f} KB")
print("\nReady for IEEE VIS 2027 VAST submission!")
