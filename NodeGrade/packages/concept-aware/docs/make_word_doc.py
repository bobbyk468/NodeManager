"""Generate problem_statement_professor_review.docx with hyperlinked TOC."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
LIGHT_BLUE = RGBColor(0x44, 0x72, 0xC4)
GREY       = RGBColor(0x70, 0x70, 0x70)
LINK_COLOR = RGBColor(0x1F, 0x49, 0x7D)

TAB_RIGHT_TWIPS = "8640"   # right-aligned tab at ~6 in from margin

# low-level helpers

def fnt(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def sp(p, before=0, after=6):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)

def _add_dot_tab(pPr):
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"),    "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"),    TAB_RIGHT_TWIPS)
    tabs.append(tab)
    pPr.append(tabs)

def _add_bookmark(p, bm_id, bm_name):
    bk_s = OxmlElement("w:bookmarkStart")
    bk_s.set(qn("w:id"),   str(bm_id))
    bk_s.set(qn("w:name"), bm_name)
    bk_e = OxmlElement("w:bookmarkEnd")
    bk_e.set(qn("w:id"),   str(bm_id))
    p._p.insert(0, bk_s)
    p._p.append(bk_e)

# TOC entry

def toc_entry(doc, label, bm_name, page_num, indent_cm=0, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(indent_cm)
    _add_dot_tab(p._p.get_or_add_pPr())

    hl = OxmlElement("w:hyperlink")
    hl.set(qn("w:anchor"), bm_name)

    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle"); rStyle.set(qn("w:val"), "Hyperlink"); rPr.append(rStyle)
    color_el = OxmlElement("w:color"); color_el.set(qn("w:val"), "1F497D"); rPr.append(color_el)
    sz_el = OxmlElement("w:sz"); sz_el.set(qn("w:val"), str(int(size * 2))); rPr.append(sz_el)
    if size >= 12:
        rPr.append(OxmlElement("w:b"))
    rPr.append(OxmlElement("w:noProof"))

    r_hl = OxmlElement("w:r"); r_hl.append(rPr)
    t_hl = OxmlElement("w:t"); t_hl.set(qn("xml:space"), "preserve"); t_hl.text = label
    r_hl.append(t_hl); hl.append(r_hl)
    p._p.append(hl)

    r_tab = OxmlElement("w:r"); r_tab.append(OxmlElement("w:tab")); p._p.append(r_tab)

    r_pg = OxmlElement("w:r")
    rPr_pg = OxmlElement("w:rPr")
    sz_pg = OxmlElement("w:sz"); sz_pg.set(qn("w:val"), str(int(size * 2))); rPr_pg.append(sz_pg)
    r_pg.append(rPr_pg)
    t_pg = OxmlElement("w:t"); t_pg.text = str(page_num); r_pg.append(t_pg)
    p._p.append(r_pg)

    return p

def add_rule(doc):
    p = doc.add_paragraph(); sp(p, 4, 4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    for k, v in [("w:val","single"),("w:sz","6"),("w:space","1"),("w:color","BFBFBF")]:
        bot.set(qn(k), v)
    pBdr.append(bot); pPr.append(pBdr)

_bm_counter = [0]

def h1(doc, n, title, bm_name):
    p = doc.add_heading(f"{n}.  {title}", level=1); sp(p, 20, 6)
    for r in p.runs:
        r.font.color.rgb = DARK_BLUE; r.font.size = Pt(14); r.font.bold = True
    _add_bookmark(p, _bm_counter[0], bm_name); _bm_counter[0] += 1
    return p

def h2(doc, title, bm_name):
    p = doc.add_heading(title, level=2); sp(p, 12, 4)
    for r in p.runs:
        r.font.color.rgb = MID_BLUE; r.font.size = Pt(12); r.font.bold = True
    _add_bookmark(p, _bm_counter[0], bm_name); _bm_counter[0] += 1
    return p

def body(doc, text, bolds=None):
    p = doc.add_paragraph(); sp(p, 0, 8); p.paragraph_format.line_spacing = Pt(14)
    if not bolds:
        r = p.add_run(text); fnt(r)
    else:
        rem = text
        for span in bolds:
            i = rem.find(span)
            if i < 0: continue
            if i > 0: r = p.add_run(rem[:i]); fnt(r)
            r = p.add_run(span); fnt(r, bold=True)
            rem = rem[i+len(span):]
        if rem: r = p.add_run(rem); fnt(r)

def qt(doc, text):
    p = doc.add_paragraph(style="Quote"); sp(p, 6, 6)
    p.paragraph_format.left_indent = Cm(1.2)
    r = p.add_run(text); fnt(r, italic=True)

def blt(doc, label, text):
    p = doc.add_paragraph(style="List Bullet"); sp(p, 2, 5)
    p.paragraph_format.left_indent = Cm(0.8)
    r = p.add_run(label); fnt(r, bold=True)
    r = p.add_run(text);  fnt(r)

def stg(doc, label, text):
    p = doc.add_paragraph(); sp(p, 6, 6); p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(label + "  "); fnt(r, bold=True, color=DARK_BLUE)
    r = p.add_run(text); fnt(r)

# BUILD DOCUMENT

doc = Document()
sec = doc.sections[0]
sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
sec.left_margin=Cm(3.0); sec.right_margin=Cm(2.5)

# cover
for txt, sz, col, ital in [
    ("ConceptGrade", 22, DARK_BLUE, False),
    ("Problem Statement, Approach, and References", 14, LIGHT_BLUE, True),
    ("Prepared for Professor Review  --  April 2026", 11, GREY, False),
]:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 6, 4)
    r = p.add_run(txt); fnt(r, size=sz, bold=(sz==22), italic=ital, color=col)

doc.add_paragraph(); add_rule(doc); doc.add_paragraph()

# TOC heading
p = doc.add_paragraph(); sp(p, 6, 10)
r = p.add_run("Table of Contents"); fnt(r, size=13, bold=True, color=DARK_BLUE)

# TOC entries -- page numbers calibrated to ~9-page rendered document
TOC = [
    ("1.  Problem Statement",                         "sec1",  2, 0,   12),
    ("2.  Research Question",                         "sec2",  3, 0,   12),
    ("3.  Proposed Approach: ConceptGrade",           "sec3",  3, 0,   12),
    ("     3.1  Core Idea",                           "sec31", 3, 0.5, 11),
    ("     3.2  The Grading Pipeline",                "sec32", 4, 0.5, 11),
    ("     3.3  What ConceptGrade Tells the Instructor", "sec33", 5, 0.5, 11),
    ("4.  Datasets",                                  "sec4",  5, 0,   12),
    ("5.  Results",                                   "sec5",  6, 0,   12),
    ("     5.1  Main Findings",                       "sec51", 6, 0.5, 11),
    ("     5.2  Component Ablation",                  "sec52", 7, 0.5, 11),
    ("     5.3  Domain Boundary Condition",           "sec53", 7, 0.5, 11),
    ("6.  Discussion and Future Work",                "sec6",  8, 0,   12),
    ("     6.1  From Scoring to Visual Analytics",    "sec61", 8, 0.5, 11),
    ("     6.2  Practical Deployment",                "sec62", 8, 0.5, 11),
    ("7.  References",                                "sec7",  9, 0,   12),
]
for label, bm, pg, indent, size in TOC:
    toc_entry(doc, label, bm, pg, indent_cm=indent, size=size)

doc.add_page_break()

# 1. PROBLEM STATEMENT
h1(doc, 1, "Problem Statement", "sec1")

body(doc,
    "Grading short written answers is one of the most demanding parts of teaching, "
    "particularly in technical subjects like Computer Science. When a class has hundreds "
    "of students, each writing a few sentences in response to questions like 'What is a "
    "linked list?' or 'How does backpropagation work?', the sheer volume of grading "
    "becomes overwhelming. Beyond the time cost, manual grading is also inconsistent -- "
    "two instructors reading the same answer often assign slightly different scores.")

body(doc,
    "Automated Short Answer Grading (ASAG) has been studied for over two decades as a "
    "solution to this problem. The goal is to build a system that reads a student's "
    "free-text answer and predicts the score a human expert would give. While meaningful "
    "progress has been made, existing approaches all share two fundamental weaknesses: "
    "they measure how similar an answer looks to the model answer rather than whether the "
    "student actually understood the underlying concepts, and they produce a number -- but "
    "nothing more. An instructor who receives a score of 2.5 out of 5 from an automated "
    "system cannot tell whether the student understood half the material well, understood "
    "everything superficially, or missed one specific concept entirely.",
    bolds=["Automated Short Answer Grading (ASAG)"])

body(doc, "The three main families of existing approaches each fail in a different but related way.")

blt(doc, "Lexical Overlap Methods (e.g., TF-IDF):  ",
    "These compare words in the student's answer against the reference. A student who "
    "memorizes the right vocabulary scores just as well as one who genuinely understands "
    "the concept. A student who explains an idea correctly using different words gets "
    "penalized unfairly.")

blt(doc, "Neural Embedding Methods (e.g., BERT, sentence transformers):  ",
    "These convert answers into vectors that capture meaning rather than surface words, "
    "handling paraphrases better. However, they still measure overall textual similarity "
    "and have no understanding of which specific concepts a student covered, which were "
    "missed, or whether the understanding is shallow or deep.")

blt(doc, "Large Language Model (LLM) Zero-Shot Grading:  ",
    "Asking models like GPT-4 or Gemini to read a student's answer and assign a score "
    "is surprisingly capable. However, LLMs suffer from cognitive-level calibration bias "
    "-- they assign high scores to fluent, confident-sounding answers regardless of "
    "whether the required concepts are actually covered. In our measurements, the pure "
    "LLM baseline consistently overestimates scores by an average of +0.19 points on a "
    "0-5 scale compared to expert human graders.")

body(doc,
    "None of these approaches can answer the questions that matter most to an instructor: "
    "Which concepts did the student correctly demonstrate? Which expected concepts are "
    "missing? Is the student's understanding shallow or deep? Does the student have "
    "misconceptions? And -- equally important -- can these answers be presented to the "
    "instructor in a form they can actually act on? These are the questions our system, "
    "ConceptGrade, is designed to answer.",
    bolds=["ConceptGrade"])

# 2. RESEARCH QUESTION
h1(doc, 2, "Research Question", "sec2")
qt(doc,
    "Can we improve both the accuracy and interpretability of automated short answer "
    "grading by explicitly checking whether a student's answer covers the expected "
    "concepts in a domain knowledge graph -- and can the resulting structured evidence "
    "be rendered as actionable visual analytics for instructors?")

body(doc,
    "This question has two parts that are inseparable in practice. Accurate scoring "
    "without explanation provides limited value to an educator. Explanations without "
    "accuracy lose credibility. ConceptGrade pursues both simultaneously: the Knowledge "
    "Graph provides structured, auditable evidence for the score, and the visualization "
    "layer converts that evidence into instructor-facing dashboards -- concept coverage "
    "heatmaps, misconception maps, Bloom's-level distributions, and student radar charts "
    "-- that make the diagnostic information actionable at the classroom scale.")

# 3. PROPOSED APPROACH
h1(doc, 3, "Proposed Approach: ConceptGrade", "sec3")
h2(doc, "3.1  Core Idea", "sec31")

body(doc,
    "The key insight behind ConceptGrade is that expert grading is fundamentally about "
    "concepts, not text. When a professor reads a student's answer, they are mentally "
    "checking whether the student has demonstrated understanding of specific ideas -- and "
    "whether those ideas are connected correctly. We can make this process explicit by "
    "representing the expert's expected knowledge as a structured Knowledge Graph (KG) "
    "and then checking how well the student's answer covers that graph.",
    bolds=["Knowledge Graph (KG)"])

body(doc,
    "A knowledge graph is a structured representation of domain knowledge as a set of "
    "concepts (nodes) and the relationships between them (edges). For the question "
    "'What is a stack?', an expert KG would contain concepts like stack, LIFO_principle, "
    "push_operation, and pop_operation, with relationships such as 'stack implements LIFO' "
    "and 'push_operation has_complexity O(1)'. A student who correctly explains all of "
    "these concepts and their connections would earn a high score. A student who only "
    "writes 'a stack stores data' covers almost none of it.")

body(doc,
    "This structured representation does something that no existing ASAG method does: "
    "it produces a named, auditable record of what the student understood and what they "
    "did not. This record is the foundation for both more accurate scoring and for the "
    "visualization layer described in Section 3.3.")

h2(doc, "3.2  The Grading Pipeline", "sec32")
body(doc, "ConceptGrade processes each grading task through four stages.")

stg(doc, "Stage 0 -- Knowledge Graph Construction.",
    "For each question, we automatically generate a domain knowledge graph using Gemini. "
    "The model reads the question and the reference answer and produces a structured JSON "
    "graph of expected concepts and their relationships. This is a one-time step per "
    "question -- once generated, the KG is stored and reused for all student answers. "
    "For the Mohler CS dataset, we designed the KG manually (101 concepts, 151 "
    "relationships) as a gold standard. For DigiKlausur and Kaggle ASAG, the KG was "
    "generated automatically, demonstrating that the pipeline scales to new domains "
    "without manual effort.")

stg(doc, "Stage 1 -- KG Feature Extraction.",
    "For each student answer, we compute concept coverage using two techniques: TF-IDF "
    "cosine similarity for exact keyword matching, and the all-MiniLM-L6-v2 "
    "sentence-transformer model for meaning-based matching that handles paraphrases. "
    "We also compute causal chain coverage and classify each answer on Bloom's taxonomy "
    "(Remember to Evaluate) and the SOLO taxonomy (Unistructural to Extended Abstract).")

stg(doc, "Stage 2 -- Scoring.",
    "We run two completely separate scoring passes using Gemini. The first is a pure LLM "
    "baseline (C_LLM) with no KG information. The second is our ConceptGrade system "
    "(C5_fix) which receives the structured KG evidence: matched concepts, coverage "
    "percentage, and the Bloom's/SOLO level. For everyday-vocabulary domains, we "
    "additionally apply an LLM-as-Judge technique: we show all expected concepts with "
    "their full descriptions and ask the model to verify each as TRUE (correctly "
    "demonstrated) or FALSE (mentioned but not understood) before scoring. This prevents "
    "bag-of-words inflation where common words like 'energy' or 'oxygen' trigger "
    "false concept matches.")

stg(doc, "Stage 3 -- Metric Computation and Visual Analytics Output.",
    "We compare scores against expert grades using MAE, Pearson correlation, Spearman "
    "rank correlation, QWK, and the Wilcoxon signed-rank test. Alongside these metrics, "
    "the pipeline produces structured JSON visualization specifications for an instructor "
    "dashboard: concept coverage charts, misconception heatmaps, student radar charts, "
    "Bloom's and SOLO distribution plots, and concept co-occurrence matrices.")

h2(doc, "3.3  What ConceptGrade Tells the Instructor", "sec33")

body(doc,
    "The distinction between ConceptGrade and all prior ASAG work is not only in the "
    "score it produces but in the structured diagnostic report that accompanies every "
    "score. For a class of 100 students answering the same question, ConceptGrade can "
    "produce: a ranked list of concepts most frequently missed; a distribution of "
    "Bloom's levels showing what fraction of the class is operating at recall versus "
    "application versus synthesis; a misconception map identifying systematic errors "
    "(e.g., students who understand the push operation but consistently misattribute the "
    "complexity of pop); and a per-student radar chart summarizing their coverage profile "
    "across five dimensions.")

body(doc,
    "These outputs transform grading from a scoring act into a diagnostic tool -- one "
    "that can inform which topics to revisit in the next lecture, which students need "
    "targeted feedback, and which parts of the question are ambiguous enough to generate "
    "systematic misunderstanding. The visualization backend (visualization/renderer.py) "
    "produces seven chart types as self-contained JSON rendering specifications, ready "
    "to be consumed by an interactive D3.js or Plotly frontend.",
    bolds=["visualization/renderer.py"])

# 4. DATASETS
h1(doc, 4, "Datasets", "sec4")
body(doc,
    "We evaluated ConceptGrade on three datasets spanning different academic domains and "
    "difficulty levels, to test whether the approach generalizes beyond the domain it was "
    "originally designed for.")

blt(doc, "Mohler et al. (2011) -- Computer Science Data Structures:  ",
    "The standard benchmark for ASAG research. 120 student answers across 10 questions "
    "on linked lists, stacks, queues, BSTs, hash tables, sorting, and Big-O notation. "
    "Scores range 0-5, assigned by expert human graders. Hand-crafted KG used.")

blt(doc, "DigiKlausur -- Neural Networks:  ",
    "646 student answers to 17 questions on neural network fundamentals, including "
    "backpropagation, activation functions, CNNs, and RNNs. Scores: 0, 2.5, or 5 "
    "(coarse rubric). Automatically generated KGs for all 17 questions.")

blt(doc, "Kaggle ASAG -- Elementary Science:  ",
    "473 student answers to 150 questions at the K-5 level, covering photosynthesis, "
    "cellular respiration, and ecosystems. Scores range 0-3. Automatically generated KGs "
    "for all 150 questions.")

# 5. RESULTS
h1(doc, 5, "Results", "sec5")
h2(doc, "5.1  Main Findings", "sec51")

body(doc,
    "ConceptGrade reduces grading error compared to the pure LLM baseline on all "
    "three datasets tested.")

body(doc,
    "On the Mohler CS dataset, ConceptGrade reduces MAE by 32.4% (0.3300 to 0.2229), "
    "with a statistically significant Wilcoxon p-value of 0.0013. It wins on MAE in "
    "8 out of 10 individual questions. The 95% confidence intervals are non-overlapping: "
    "ConceptGrade [0.179, 0.269] vs. LLM baseline [0.273, 0.390].",
    bolds=["32.4%", "p-value of 0.0013"])

body(doc,
    "On the DigiKlausur Neural Networks dataset, ConceptGrade reduces MAE by 4.9% "
    "(1.1842 to 1.1262), with a Wilcoxon p-value of 0.049.",
    bolds=["4.9%", "p-value of 0.049"])

body(doc,
    "On the Kaggle ASAG Science dataset, ConceptGrade reduces MAE by 3.2% "
    "(1.2082 to 1.1691). This improvement is directionally consistent but does not "
    "reach statistical significance independently (p = 0.319). This is discussed in "
    "Section 5.3.",
    bolds=["3.2%"])

body(doc,
    "Combining all three datasets, a Fisher combined test across all three p-values "
    "yields p = 0.0014, confirming that the overall evidence for ConceptGrade's advantage "
    "is highly significant across 1,239 student answers spanning Computer Science, "
    "Neural Networks, and Elementary Science.",
    bolds=["p = 0.0014"])

h2(doc, "5.2  Component Ablation (Mohler Dataset)", "sec52")
body(doc,
    "To understand which parts of the pipeline contribute most, we tested all intermediate "
    "configurations on the Mohler dataset:")

for lbl, txt in [
    ("C0 (reference answer only, no student answer):  ",
     "MAE = 1.7113 -- the KG alone is useless without reading the student's answer."),
    ("C1 (KG evidence only, no reference answer):  ",
     "MAE = 0.7405 -- better, but missing critical rubric context."),
    ("C1_fix (KG + student answer, no reference answer):  ",
     "MAE = 0.3458 -- close to but slightly worse than the LLM baseline."),
    ("C_LLM (full LLM baseline, no KG):  ", "MAE = 0.3300"),
    ("CoT baseline (chain-of-thought prompting):  ", "MAE = 0.2208"),
    ("ConceptGrade / C5_fix (full pipeline):  ", "MAE = 0.2229"),
]:
    blt(doc, lbl, txt)

body(doc,
    "ConceptGrade essentially matches the CoT baseline in raw accuracy. The key "
    "difference is interpretability: ConceptGrade tells you which concepts the student "
    "covered and which are missing, while CoT produces only a score and a one-line "
    "justification. The structured concept coverage data is precisely what feeds the "
    "visualization renderer.",
    bolds=["interpretability"])

h2(doc, "5.3  Domain Boundary Condition", "sec53")

body(doc,
    "The smaller benefit on Kaggle ASAG is not a failure -- it is a theoretically "
    "important finding. Elementary science questions use everyday vocabulary: energy, "
    "water, oxygen, plants. A student can write a fluent sentence using all these words "
    "while explaining the concept entirely incorrectly. Concept keyword presence becomes "
    "a weak signal of understanding when vocabulary is common to everyday language.")

body(doc,
    "By contrast, Computer Science and Neural Networks use specialized vocabulary -- "
    "backpropagation, bipartite graph, O(n log n) -- where using a term correctly is "
    "itself strong evidence of understanding. This points to a general principle: the "
    "benefit of Knowledge Graph augmentation scales with the lexical specificity of the "
    "domain. This boundary condition has direct practical implications for where "
    "ConceptGrade should be deployed, and it also affects the reliability of the "
    "visualization layer: in technical domains, concept coverage heatmaps and misconception "
    "maps are highly informative; in everyday-language domains, additional LLM-as-Judge "
    "verification is needed to prevent false-positive matches from inflating the visual "
    "diagnosis.",
    bolds=["the benefit of Knowledge Graph augmentation scales with the lexical specificity of the domain."])

# 6. DISCUSSION AND FUTURE WORK
h1(doc, 6, "Discussion and Future Work", "sec6")
h2(doc, "6.1  From Scoring to Explainable Visual Analytics", "sec61")

body(doc,
    "ConceptGrade's current design already produces substantially more than a score. "
    "Every graded answer generates a structured diagnostic record: which concepts were "
    "matched, at what Bloom's and SOLO level, and whether the student's reasoning shows "
    "causal integration or isolated recall. At the class level, these records aggregate "
    "into maps of collective understanding and misconception.")

body(doc,
    "The next natural step for this work is building the interactive instructor-facing "
    "layer that consumes these structured outputs. The backend visualization engine is "
    "complete -- it generates JSON rendering specifications for seven visualization types: "
    "concept coverage charts, misconception heatmaps, student radar charts, co-occurrence "
    "matrices, Bloom's and SOLO distribution plots. The missing piece is a D3.js or "
    "Plotly frontend that renders these specifications interactively, and a user study "
    "with actual instructors to validate that the diagnostic outputs are useful and "
    "actionable in classroom practice.")

body(doc,
    "This combination -- KG-grounded accuracy, structured diagnostic evidence, and "
    "interactive visual exploration -- positions ConceptGrade as a Visual Analytics "
    "system for explainable AI grading, a framing that opens up a distinct research "
    "direction from the crowded ASAG benchmarking literature. The VAST (Visual Analytics "
    "Science and Technology) track of IEEE VIS represents a fitting venue for presenting "
    "this work once the interactive frontend and educator study are complete.",
    bolds=["IEEE VIS"])

h2(doc, "6.2  Practical Deployment Considerations", "sec62")

body(doc,
    "Based on the multi-dataset evaluation, the clearest deployment targets for "
    "ConceptGrade are university-level courses in technical domains -- Computer Science, "
    "engineering, and applied mathematics -- where specialized vocabulary makes KG "
    "matching highly reliable. For K-12 science or humanities courses, the LLM-as-Judge "
    "enhancement mitigates the vocabulary specificity problem, but further "
    "domain-specific calibration would be needed before production deployment.")

body(doc,
    "The automatic KG generation pipeline already scales to new domains without manual "
    "KG design -- the DigiKlausur and Kaggle datasets both used fully automated KG "
    "construction. A course instructor who wanted to deploy ConceptGrade on their own "
    "question bank would need only to provide the questions and reference answers; the "
    "system would build the KGs automatically and begin producing diagnostic outputs "
    "with no additional configuration.")

# 7. REFERENCES
h1(doc, 7, "References", "sec7")

REFS = [
    ("Mohler, M., & Mihalcea, R. (2009).",
     "Text-to-text semantic similarity for automatic short answer grading. EACL 2009, pp. 567-575."),
    ("Mohler, M., Bunescu, R., & Mihalcea, R. (2011).",
     "Learning to grade short answer questions using semantic similarity measures and dependency "
     "graph alignments. ACL 2011, pp. 752-762. [Primary benchmark used in this work]"),
    ("Sultan, M. A., Salazar, C., & Sumner, T. (2016).",
     "Fast and easy short answer grading with high accuracy. NAACL-HLT 2016, pp. 1-11."),
    ("Dzikovska, M. et al. (2013).",
     "SemEval-2013 Task 7: The Joint Student Response Analysis and 8th Recognizing Textual "
     "Entailment Challenge. SemEval 2013, pp. 263-274."),
    ("Devlin, J., Chang, M.-W., Lee, K., & Toutanova, K. (2019).",
     "BERT: Pre-training of deep bidirectional transformers for language understanding. "
     "NAACL-HLT 2019, pp. 4171-4186."),
    ("Sung, T., Nain, N., & Sharma, H. K. (2023).",
     "Pre-trained language model-based automatic short answer grading for CS education. "
     "IEEE Access, vol. 11, pp. 31,542-31,553."),
    ("Bhandari, P. et al. (2023).",
     "Can ChatGPT replace human evaluators? An empirical study of automatic short answer "
     "grading. arXiv:2308.12505."),
    ("Gao, H. et al. (2024).",
     "Automated short answer grading using large language models: A zero-shot and few-shot "
     "exploration. IEEE ICALT 2024."),
    ("Emirtekin, E., & Ozarslan, Y. (2025).",
     "Automated classification of student responses using Bloom's taxonomy with large language "
     "models. Computers & Education: Artificial Intelligence, vol. 8."),
    ("Bloom, B. S. et al. (1956).",
     "Taxonomy of Educational Objectives: Handbook I: Cognitive Domain. New York: David McKay."),
    ("Anderson, L. W., & Krathwohl, D. R. (Eds.) (2001).",
     "A Taxonomy for Learning, Teaching, and Assessing: A Revision of Bloom's Taxonomy. Longman."),
    ("Biggs, J. B., & Collis, K. F. (1982).",
     "Evaluating the Quality of Learning: The SOLO Taxonomy. New York: Academic Press."),
    ("Corbett, A. T., & Anderson, J. R. (1994).",
     "Knowledge tracing: Modeling the acquisition of procedural knowledge. "
     "User Modeling and User-Adapted Interaction, vol. 4, no. 4, pp. 253-278."),
    ("Pan, C., Li, N., Rusakov, M., & Faloutsos, C. (2017).",
     "Prerequisite relation learning for concepts in MOOCs. ACL 2017, pp. 1447-1456."),
    ("Leacock, C., & Chodorow, M. (2003).",
     "C-rater: Automated scoring of short-answer questions. "
     "Computers in Human Behavior, vol. 19, no. 4, pp. 491-508."),
    ("Landis, J. R., & Koch, G. G. (1977).",
     "The measurement of observer agreement for categorical data. "
     "Biometrics, vol. 33, no. 1, pp. 159-174."),
    ("Keim, D., Kohlhammer, J., Ellis, G., & Mansmann, F. (Eds.) (2010).",
     "Mastering the Information Age: Solving Problems with Visual Analytics. "
     "Goslar: Eurographics Association."),
    ("Amershi, S. et al. (2019).",
     "Software engineering for machine learning: A case study. "
     "Proceedings of ICSE-SEIP 2019, pp. 291-300."),
]

for i, (author, detail) in enumerate(REFS, 1):
    p = doc.add_paragraph(); sp(p, 1, 5)
    p.paragraph_format.left_indent       = Cm(0.9)
    p.paragraph_format.first_line_indent = Cm(-0.9)
    r = p.add_run(f"[{i}]  {author}  "); fnt(r, size=10.5, bold=True)
    r = p.add_run(detail); fnt(r, size=10.5)

# save
out = "problem_statement_professor_review.docx"
doc.save(out)
print(f"Saved -> {out}")
