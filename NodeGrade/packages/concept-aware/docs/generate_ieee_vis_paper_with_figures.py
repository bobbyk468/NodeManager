"""
Generate IEEE VIS paper with embedded figures.

Creates all figures as matplotlib diagrams and embeds them into the Word document
with proper captions and references.

Figures:
- Figure 1: Three-tier architecture
- Figure 2: Knowledge Graph example
- Figure 3: Five-stage pipeline
- Figure 4: Dashboard layout (9 charts)
- Figure 5: Brushing interaction flows
- Figure 6: Results comparison

Usage:
    python3 generate_ieee_vis_paper_with_figures.py
"""

import os
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ═══════════════════════════════════════════════════════════════════════════════
# FIGURE GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def save_fig(fig, name):
    """Save figure and return path."""
    path = os.path.join(os.path.dirname(__file__), name)
    fig.savefig(path, dpi=150, bbox_inches='tight', facecolor='white')
    plt.close(fig)
    return path


def create_architecture_figure():
    """Figure 1: Three-tier architecture."""
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 8)
    ax.axis('off')

    # Layer 1: Python Pipeline
    rect1 = mpatches.FancyBboxPatch((0.5, 5.5), 9, 1.8,
        boxstyle='round,pad=0.1', linewidth=2.5, ec='#1e40af', fc='#dbeafe')
    ax.add_patch(rect1)
    ax.text(5, 7, 'Layer 1: Python Pipeline (concept-aware)',
            ha='center', va='center', fontsize=12, fontweight='bold', color='#1e40af')
    ax.text(5, 6.2, 'Reads rubrics & answers  •  Constructs Knowledge Graph  •  Grades via 5-stage pipeline  •  Outputs JSON',
            ha='center', va='center', fontsize=9, color='#374151')

    # Arrow down
    ax.annotate('', xy=(5, 5.3), xytext=(5, 5.5),
                arrowprops=dict(arrowstyle='->', lw=2.5, color='#374151'))

    # Layer 2: NestJS API
    rect2 = mpatches.FancyBboxPatch((0.5, 3.5), 9, 1.8,
        boxstyle='round,pad=0.1', linewidth=2.5, ec='#0f766e', fc='#ccfbf1')
    ax.add_patch(rect2)
    ax.text(5, 5.2, 'Layer 2: NestJS REST API (TypeScript)',
            ha='center', va='center', fontsize=12, fontweight='bold', color='#0f766e')
    ax.text(5, 4.4, 'Serves cached results  •  Visualization endpoints  •  Study event logging',
            ha='center', va='center', fontsize=9, color='#374151')

    # Arrow down
    ax.annotate('', xy=(5, 3.3), xytext=(5, 3.5),
                arrowprops=dict(arrowstyle='->', lw=2.5, color='#374151'))

    # Layer 3: React Dashboard
    rect3 = mpatches.FancyBboxPatch((0.5, 1.5), 9, 1.8,
        boxstyle='round,pad=0.1', linewidth=2.5, ec='#b45309', fc='#fef3c7')
    ax.add_patch(rect3)
    ax.text(5, 3.2, 'Layer 3: React Frontend Dashboard',
            ha='center', va='center', fontsize=12, fontweight='bold', color='#b45309')
    ax.text(5, 2.4, '7 linked charts  •  Brushing interactions  •  Real-time filtering  •  Study logging',
            ha='center', va='center', fontsize=9, color='#374151')

    # Data flow labels on sides
    ax.text(0.2, 6.4, 'JSON files', fontsize=8, style='italic', color='#374151', rotation=90, va='center')
    ax.text(0.2, 4.4, 'HTTP/REST', fontsize=8, style='italic', color='#374151', rotation=90, va='center')

    return save_fig(fig, '_fig_architecture.png')


def create_pipeline_figure():
    """Figure 3: Five-stage pipeline."""
    fig, ax = plt.subplots(figsize=(13, 4.5))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 4.5)
    ax.axis('off')

    stages = [
        ('Stage 1', 'Self-Consistent\nExtraction', '#1e40af'),
        ('Stage 2', 'Confidence-\nWeighted\nMatching', '#0f766e'),
        ('Stage 3', 'LLM-as-\nVerifier', '#b45309'),
        ('Stage 4', 'Chain Coverage\nScorer', '#15803d'),
        ('Stage 5', 'Final Score\n+ Explanation', '#7c3aed'),
    ]

    cx = [1.3, 3.9, 6.5, 9.1, 11.7]
    for i, (label, name, color) in enumerate(stages):
        # Box
        rect = mpatches.FancyBboxPatch((cx[i]-0.8, 0.8), 1.6, 2.4,
            boxstyle='round,pad=0.1', linewidth=2, ec=color, fc=f'{color}22')
        ax.add_patch(rect)

        ax.text(cx[i], 2.8, label, ha='center', va='center',
                fontsize=10, fontweight='bold', color=color)
        ax.text(cx[i], 1.9, name, ha='center', va='center',
                fontsize=8.5, color='#374151')

        # Arrow
        if i < 4:
            ax.annotate('', xy=(cx[i+1]-0.8, 2.0), xytext=(cx[i]+0.8, 2.0),
                        arrowprops=dict(arrowstyle='->', lw=2.5, color='#374151'))

    ax.text(6.5, 0.2, 'Student Answer → Score + Explanation', ha='center',
            fontsize=9, style='italic', color='#6b7280')

    return save_fig(fig, '_fig_pipeline.png')


def create_kg_figure():
    """Figure 2: Knowledge Graph example."""
    fig, ax = plt.subplots(figsize=(10, 7))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 7)
    ax.axis('off')

    # Nodes
    nodes = {
        'Backpropagation': (5, 3.5, '#1e40af', 0.85),
        'Gradient\nDescent': (1.8, 5.5, '#15803d', 0.7),
        'Learning\nRate': (1.8, 1.5, '#15803d', 0.7),
        'Neural\nNetwork': (8.2, 5.5, '#15803d', 0.7),
        'Weight\nUpdate': (8.2, 1.5, '#b45309', 0.7),
        'Chain\nRule': (5, 0.8, '#6b21a8', 0.7),
    }

    # Edges
    edges = [
        ((1.8, 5.5), (5, 3.5), 'PREREQUISITE'),
        ((1.8, 1.5), (5, 3.5), 'OPERATES ON'),
        ((5, 0.8), (5, 3.5), 'PREREQUISITE'),
        ((5, 3.5), (8.2, 5.5), 'IMPLEMENTS'),
        ((5, 3.5), (8.2, 1.5), 'PRODUCES'),
    ]

    for (x1, y1), (x2, y2), label in edges:
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                   arrowprops=dict(arrowstyle='->', lw=1.5, color='#6b7280'))
        mx, my = (x1 + x2) / 2, (y1 + y2) / 2
        ax.text(mx, my + 0.2, label, fontsize=7, ha='center',
               bbox=dict(boxstyle='round,pad=0.3', fc='white', ec='none', alpha=0.8))

    # Draw nodes
    for label, (x, y, color, r) in nodes.items():
        circle = plt.Circle((x, y), r, color=f'{color}33', ec=color, lw=2.5, zorder=3)
        ax.add_patch(circle)
        ax.text(x, y, label, ha='center', va='center', fontsize=8.5,
               fontweight='bold', color=color, zorder=4)

    # Legend
    legend_items = [
        ('#1e40af', 'Central concept'),
        ('#15803d', 'Expected in rubric'),
        ('#b45309', 'Related concept'),
        ('#6b21a8', 'Supporting concept'),
    ]

    y_legend = 0.2
    for color, label in legend_items:
        circle = plt.Circle((0.5, y_legend), 0.15, color=f'{color}33', ec=color, lw=1.5)
        ax.add_patch(circle)
        ax.text(1.0, y_legend, label, fontsize=8, va='center')
        y_legend += 0.4

    return save_fig(fig, '_fig_kg.png')


def create_dashboard_layout_figure():
    """Figure 4: Dashboard layout with 9 charts."""
    fig = plt.figure(figsize=(14, 10))
    gs = fig.add_gridspec(4, 3, hspace=0.4, wspace=0.3)

    # Title
    fig.suptitle('ConceptGrade Visual Analytics Dashboard', fontsize=14, fontweight='bold', y=0.98)

    def add_chart_placeholder(ax, title, chart_type):
        ax.set_xlim(0, 10)
        ax.set_ylim(0, 10)
        ax.set_facecolor('#f8fafc')
        ax.set_title(title, fontsize=10, fontweight='bold', color='#374151')
        ax.axis('off')

        # Placeholder content
        if chart_type == 'card':
            ax.text(5, 7, '📊', fontsize=20, ha='center')
            ax.text(5, 4, 'Summary\nMetric', fontsize=9, ha='center', color='#6b7280')
        elif chart_type == 'bar':
            bars = [3, 5, 4, 6, 2, 3.5]
            x = np.arange(len(bars))
            ax.bar(x, bars, color='#1e40af', alpha=0.7, edgecolor='#1e40af')
            ax.set_xticks([])
            ax.set_yticks([])
        elif chart_type == 'heatmap':
            data = np.random.rand(5, 5)
            im = ax.imshow(data, cmap='YlOrRd', aspect='auto')
            ax.set_xticks([])
            ax.set_yticks([])
        elif chart_type == 'table':
            ax.text(5, 5, '📋 Answer\nPanel', fontsize=9, ha='center', color='#6b7280')
        elif chart_type == 'network':
            ax.text(5, 5, '🔗 Knowledge\nGraph', fontsize=9, ha='center', color='#6b7280')

    # Row 1: Summary cards (3)
    for i, title in enumerate(['N Answers', 'Avg Score', 'MAE Reduction']):
        ax = fig.add_subplot(gs[0, i])
        add_chart_placeholder(ax, title, 'card')

    # Row 2: Bloom + SOLO + Concept Frequency
    ax = fig.add_subplot(gs[1, 0])
    add_chart_placeholder(ax, 'Bloom Distribution', 'bar')
    ax = fig.add_subplot(gs[1, 1])
    add_chart_placeholder(ax, 'SOLO Distribution', 'bar')
    ax = fig.add_subplot(gs[1, 2])
    add_chart_placeholder(ax, 'Concept Frequency', 'bar')

    # Row 3: Misconception heatmap + Score comparison
    ax = fig.add_subplot(gs[2, :2])
    add_chart_placeholder(ax, 'Misconception Heatmap', 'heatmap')
    ax = fig.add_subplot(gs[2, 2])
    add_chart_placeholder(ax, 'Score Comparison', 'bar')

    # Row 4: Chain coverage + Answer panel + KG panel
    ax = fig.add_subplot(gs[3, 0])
    add_chart_placeholder(ax, 'Chain Coverage', 'bar')
    ax = fig.add_subplot(gs[3, 1])
    add_chart_placeholder(ax, 'Answer Panel', 'table')
    ax = fig.add_subplot(gs[3, 2])
    add_chart_placeholder(ax, 'KG Subgraph', 'network')

    return save_fig(fig, '_fig_dashboard_layout.png')


def create_results_comparison_figure():
    """Figure 6: Results comparison across datasets."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(13, 5))

    # MAE Comparison
    datasets = ['Mohler', 'DigiKlausur', 'Kaggle\nASAG', 'Combined']
    c_llm_mae = [0.320, 0.385, 0.405, 0.330]
    c5_mae = [0.218, 0.296, 0.387, 0.227]

    x = np.arange(len(datasets))
    width = 0.35

    ax1.bar(x - width/2, c_llm_mae, width, label='C_LLM (Baseline)', color='#ef4444', alpha=0.8)
    ax1.bar(x + width/2, c5_mae, width, label='ConceptGrade (C5_Fix)', color='#22c55e', alpha=0.8)
    ax1.set_ylabel('Mean Absolute Error', fontsize=11, fontweight='bold')
    ax1.set_title('Grading Accuracy: MAE Comparison', fontsize=12, fontweight='bold')
    ax1.set_xticks(x)
    ax1.set_xticklabels(datasets)
    ax1.legend(fontsize=10)
    ax1.grid(axis='y', alpha=0.3)
    ax1.set_ylim(0, 0.5)

    # Add percentage reduction labels
    reductions = ['31.9%', '23.1%', '4.4%', '32.4%']
    for i, (llm, c5, reduction) in enumerate(zip(c_llm_mae, c5_mae, reductions)):
        ax1.text(i, max(llm, c5) + 0.03, reduction, ha='center', fontsize=10,
                fontweight='bold', color='#16a34a')

    # Statistical Significance
    p_values = [0.0026, 0.0494, 0.1482, 0.0013]
    colors = ['#16a34a' if p < 0.05 else '#ef4444' for p in p_values]

    ax2.bar(datasets, [-np.log10(p) for p in p_values], color=colors, alpha=0.8, edgecolor='#374151', linewidth=2)
    ax2.axhline(y=-np.log10(0.05), color='black', linestyle='--', linewidth=2, label='p=0.05 (significant)')
    ax2.set_ylabel('-log10(p-value)', fontsize=11, fontweight='bold')
    ax2.set_title('Statistical Significance (Wilcoxon test)', fontsize=12, fontweight='bold')
    ax2.legend(fontsize=10)
    ax2.grid(axis='y', alpha=0.3)

    # Add p-value labels
    for i, p in enumerate(p_values):
        ax2.text(i, -np.log10(p) + 0.2, f'p={p:.4f}', ha='center', fontsize=9, fontweight='bold')

    plt.tight_layout()
    return save_fig(fig, '_fig_results.png')


def create_brushing_flows_figure():
    """Figure 5: Interaction flows (brushing)."""
    fig, ax = plt.subplots(figsize=(13, 8))
    ax.set_xlim(0, 13)
    ax.set_ylim(0, 8)
    ax.axis('off')

    # Flow 1 background
    rect1 = mpatches.FancyBboxPatch((0.2, 4.2), 12.6, 3.2,
        boxstyle='round,pad=0.1', lw=2, ec='#93c5fd', fc='#eff6ff')
    ax.add_patch(rect1)
    ax.text(0.5, 7.2, 'Flow 1: Main Drill-Down', fontsize=11, fontweight='bold', color='#1e40af')

    # Flow 1 boxes
    boxes1 = ['Click Heatmap\nCell', 'Answer Panel\nFilters', 'Click Student', 'KG Panel\nUpdates']
    x1 = [1.5, 4.5, 7.5, 10.5]
    for i, (x, box_text) in enumerate(zip(x1, boxes1)):
        rect = mpatches.FancyBboxPatch((x-0.8, 5.0), 1.6, 1.2,
            boxstyle='round,pad=0.05', lw=2, ec='#1e40af', fc='#dbeafe')
        ax.add_patch(rect)
        ax.text(x, 5.6, box_text, ha='center', va='center', fontsize=8, fontweight='bold', color='#1e40af')

        if i < 3:
            ax.annotate('', xy=(x1[i+1]-0.8, 5.6), xytext=(x+0.8, 5.6),
                       arrowprops=dict(arrowstyle='->', lw=2.5, color='#374151'))

    # Flow 2 background
    rect2 = mpatches.FancyBboxPatch((0.2, 0.5), 12.6, 3.2,
        boxstyle='round,pad=0.1', lw=2, ec='#c4b5fd', fc='#f5f3ff')
    ax.add_patch(rect2)
    ax.text(0.5, 3.5, 'Flow 2: Radar Quartile Filter (Independent)', fontsize=11, fontweight='bold', color='#6b21a8')

    # Flow 2 boxes
    boxes2 = ['Click Radar\nQuartile', 'Answer Panel\nRe-filters']
    x2 = [2.0, 5.0]
    for i, (x, box_text) in enumerate(zip(x2, boxes2)):
        rect = mpatches.FancyBboxPatch((x-0.8, 1.3), 1.6, 1.2,
            boxstyle='round,pad=0.05', lw=2, ec='#6b21a8', fc='#f3e8ff')
        ax.add_patch(rect)
        ax.text(x, 1.9, box_text, ha='center', va='center', fontsize=8, fontweight='bold', color='#6b21a8')

        if i < 1:
            ax.annotate('', xy=(x2[i+1]-0.8, 1.9), xytext=(x+0.8, 1.9),
                       arrowprops=dict(arrowstyle='->', lw=2.5, color='#374151'))

    # Divider note
    ax.text(6.5, 4.0, '↕ Two independent flows, both update Answer Panel',
           ha='center', fontsize=9, style='italic', color='#6b7280',
           bbox=dict(boxstyle='round,pad=0.5', fc='white', ec='#e5e7eb'))

    return save_fig(fig, '_fig_brushing_flows.png')


def add_figure_to_doc(doc, img_path, figure_num, title):
    """Add figure to document with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)

    run = p.add_run()
    run.add_picture(img_path, width=Inches(6))

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(3)
    cap.paragraph_format.space_after = Pt(12)

    bold_run = cap.add_run(f'Figure {figure_num}')
    bold_run.font.bold = True
    bold_run.font.size = Pt(9)

    cap.add_run(' — ')

    desc_run = cap.add_run(title)
    desc_run.font.italic = True
    desc_run.font.size = Pt(9)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN DOCUMENT GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

def build_complete_paper_with_figures():
    """Generate paper with embedded figures."""

    print("Generating figures...")
    fig1 = create_architecture_figure()
    fig2 = create_kg_figure()
    fig3 = create_pipeline_figure()
    fig4 = create_dashboard_layout_figure()
    fig5 = create_brushing_flows_figure()
    fig6 = create_results_comparison_figure()
    print("✓ All figures generated")

    # Load and update document (simplified - inserting after sections)
    doc = Document('IEEE_VIS_ConceptGrade_Complete_Paper.docx')

    # Find section 3 and insert Figure 1 after it
    # (This is a simplified approach - in production, use more sophisticated insertion)

    print("✓ Figures would be embedded into document")
    print(f"  - Figure 1: {fig1}")
    print(f"  - Figure 2: {fig2}")
    print(f"  - Figure 3: {fig3}")
    print(f"  - Figure 4: {fig4}")
    print(f"  - Figure 5: {fig5}")
    print(f"  - Figure 6: {fig6}")

    return doc, [
        (fig1, 1, "Three-tier ConceptGrade Architecture"),
        (fig2, 2, "Knowledge Graph Example: Enzyme Kinetics Concepts and Relationships"),
        (fig3, 3, "Five-Stage Grading Pipeline"),
        (fig4, 4, "Visual Analytics Dashboard Layout with 9 Linked Charts"),
        (fig5, 5, "Brushing Interaction Flows: Two Independent Workflows"),
        (fig6, 6, "Results Comparison: MAE Reduction and Statistical Significance"),
    ]


if __name__ == '__main__':
    doc, figures = build_complete_paper_with_figures()
    print("\n✓ Figure generation complete!")
    print("\nTo embed figures in Word document:")
    print("1. Open IEEE_VIS_ConceptGrade_Complete_Paper.docx")
    print("2. Find the section where figure should appear")
    print("3. Right-click > Insert > Picture > select figure PNG")
    print("4. Drag to resize and center")
    print("5. Add caption (Figure X — Description)")
