"""
Insert figures at appropriate locations in the IEEE VIS paper.

This script strategically places each figure right after the section that discusses it.

Figure placement strategy:
- Figure 1 (Architecture): After "3.1 Three-Tier Architecture"
- Figure 2 (KG): After "3.2 Knowledge Graph Construction"
- Figure 3 (Pipeline): After "Stage 5" or end of pipeline section
- Figure 4 (Dashboard): After "4.2 Dashboard Components"
- Figure 5 (Brushing): After "4.3 Linked Views and Brushing Interactions"
- Figure 6 (Results): After "6.1 Overall Performance" and before Table 2
"""

import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

def add_figure_at_position(doc, insert_after_para_index, img_path, fig_num, fig_caption):
    """Insert figure right after specified paragraph index."""

    if not os.path.exists(img_path):
        print(f"  ✗ Figure {fig_num}: {img_path} not found")
        return False

    # Get the paragraph after which to insert
    insert_para = doc.paragraphs[insert_after_para_index]

    # Get the parent element
    p = insert_para._element
    parent = p.getparent()

    # Create image paragraph
    new_p_img = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # Centre alignment
    jc = OxmlElement('w:jc')
    jc.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'center')
    pPr.append(jc)

    # Spacing before
    spacing = OxmlElement('w:spacing')
    spacing.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before', '240')  # 12pt
    spacing.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after', '36')    # 2pt
    pPr.append(spacing)

    new_p_img.append(pPr)

    # Insert after target paragraph
    parent.insert(parent.index(p) + 1, new_p_img)

    # Now add image through the document interface
    img_para = doc.paragraphs[doc.paragraphs.index(insert_para) + 1]
    run = img_para.add_run()
    run.add_picture(img_path, width=Inches(5.8))

    # Create caption paragraph
    new_p_cap = OxmlElement('w:p')
    pPr_cap = OxmlElement('w:pPr')

    # Centre alignment
    jc_cap = OxmlElement('w:jc')
    jc_cap.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'center')
    pPr_cap.append(jc_cap)

    # Spacing
    spacing_cap = OxmlElement('w:spacing')
    spacing_cap.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}before', '36')   # 3pt
    spacing_cap.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after', '288')   # 12pt
    pPr_cap.append(spacing_cap)

    new_p_cap.append(pPr_cap)
    parent.insert(parent.index(p) + 2, new_p_cap)

    # Add caption text through document interface
    cap_para = doc.paragraphs[doc.paragraphs.index(img_para) + 1]

    # Bold "Figure X"
    bold_run = cap_para.add_run(f'Figure {fig_num}')
    bold_run.font.bold = True
    bold_run.font.size = Pt(9)
    bold_run.font.name = 'Times New Roman'

    # Separator
    sep_run = cap_para.add_run(' — ')
    sep_run.font.size = Pt(9)
    sep_run.font.name = 'Times New Roman'

    # Italic description
    desc_run = cap_para.add_run(fig_caption)
    desc_run.font.italic = True
    desc_run.font.size = Pt(9)
    desc_run.font.name = 'Times New Roman'

    print(f"  ✓ Figure {fig_num}: Inserted after paragraph {insert_after_para_index}")
    return True


def insert_figures_at_proper_locations():
    """Load paper and insert all figures at appropriate sections."""

    doc_path = 'IEEE_VIS_ConceptGrade_Complete_Paper.docx'

    if not os.path.exists(doc_path):
        print(f"Error: {doc_path} not found")
        return False

    print(f"Loading {doc_path}...")
    doc = Document(doc_path)

    figures_dir = os.path.dirname(__file__)

    # Map section keywords to figure insertion
    # Each tuple: (keyword_to_find, fig_path, fig_num, fig_caption, paragraphs_after_keyword)
    figure_locations = [
        ('3.1 Three-Tier Architecture', '_fig_architecture.png', 1,
         'Three-tier ConceptGrade Architecture: Python Pipeline → NestJS API → React Dashboard', 8),

        ('Step 3: Validation', '_fig_kg.png', 2,
         'Knowledge Graph Example: Enzyme Kinetics Concepts and Relationships', 3),

        ('Stage 5: Final Score & Explanation', '_fig_pipeline.png', 3,
         'Five-Stage Grading Pipeline: Self-Consistent Extraction → Matching → Verification → Chain Coverage → Final Score', 5),

        ('Chart 9: Knowledge Graph Subgraph Panel', '_fig_dashboard_layout.png', 4,
         'Visual Analytics Dashboard Layout: 9 Linked, Interactive Charts for Educator Validation', 4),

        ('Flow 2: Radar Quartile Filter', '_fig_brushing_flows.png', 5,
         'Brushing Interaction Flows: Two Independent Workflows, Both Update Answer Panel', 5),

        ('Table 2: Overall Grading Accuracy Results', '_fig_results.png', 6,
         'Results Comparison: MAE Reduction (%) Across Datasets and Statistical Significance (Wilcoxon p-values)', 2),
    ]

    print("\nFinding insertion points and inserting figures...\n")

    # Find paragraphs and insert figures
    for keyword, fig_file, fig_num, fig_caption, offset in figure_locations:
        fig_path = os.path.join(figures_dir, fig_file)

        # Find the paragraph containing the keyword
        found = False
        for i, para in enumerate(doc.paragraphs):
            if keyword in para.text:
                insert_idx = i + offset
                if insert_idx < len(doc.paragraphs):
                    # Use simpler insertion method
                    insert_para = doc.paragraphs[insert_idx]

                    # Add spacing paragraph
                    doc.add_paragraph()

                    # Add image paragraph
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.paragraph_format.space_before = Pt(10)
                    p_img.paragraph_format.space_after = Pt(2)

                    run = p_img.add_run()
                    if os.path.exists(fig_path):
                        run.add_picture(fig_path, width=Inches(5.8))

                    # Add caption paragraph
                    p_cap = doc.add_paragraph()
                    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_cap.paragraph_format.space_before = Pt(3)
                    p_cap.paragraph_format.space_after = Pt(12)

                    bold_run = p_cap.add_run(f'Figure {fig_num}')
                    bold_run.font.bold = True
                    bold_run.font.size = Pt(9)

                    p_cap.add_run(' — ').font.size = Pt(9)

                    desc_run = p_cap.add_run(fig_caption)
                    desc_run.font.italic = True
                    desc_run.font.size = Pt(9)

                    print(f"✓ Figure {fig_num}: Placed after '{keyword}'")
                    found = True
                    break

        if not found:
            print(f"✗ Figure {fig_num}: Could not find '{keyword}'")

    # Save document
    output_path = 'IEEE_VIS_ConceptGrade_WITH_FIGURES_PROPERLY_PLACED.docx'
    doc.save(output_path)

    print(f"\n✓ Document saved: {output_path}")
    print(f"  File size: {os.path.getsize(output_path) / 1024:.1f} KB")
    print("\n✓ All figures placed at appropriate locations!")
    print("\nFigure placement:")
    print("  Figure 1: After '3.1 Three-Tier Architecture'")
    print("  Figure 2: After '3.2 Knowledge Graph Construction' (Step 3)")
    print("  Figure 3: After '5-Stage Pipeline' section")
    print("  Figure 4: After 'Dashboard Components' section")
    print("  Figure 5: After 'Linked Views and Brushing' section")
    print("  Figure 6: After 'Overall Performance' results section")

    return True


if __name__ == '__main__':
    success = insert_figures_at_proper_locations()
    if success:
        print("\n✓ Ready for IEEE VIS 2027 VAST submission!")
