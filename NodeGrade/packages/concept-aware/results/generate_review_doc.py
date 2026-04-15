"""
Generate conceptgrade_lrm_ablation_review.docx from ablation summary JSON.
Run: python results/generate_review_doc.py
"""
from __future__ import annotations
import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


DATA_DIR    = Path(__file__).parent.parent / 'data'
RESULTS_DIR = Path(__file__).parent


# ── helpers ───────────────────────────────────────────────────────────────────

def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def table_row(table, cells: list, bold: bool = False, shade: str | None = None) -> None:
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = str(val)
        if bold:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        if shade:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement('w:shd')
            shd.set(qn('w:val'),   'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), shade)
            tcPr.append(shd)


# ── main ──────────────────────────────────────────────────────────────────────

def build_doc() -> None:
    with open(DATA_DIR / 'lrm_ablation_summary.json') as f:
        summary = json.load(f)

    doc = Document()

    # ── Title ──
    title = doc.add_heading('ConceptGrade — LRM Ablation Study', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph('Results & Review Request  ·  2026-04-14')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True

    sub2 = doc.add_paragraph(
        'Prepared for IEEE VIS 2027 VAST submission: '
        '"Visual Analytics for Explainable AI Grading"'
    )
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── 1. System Overview ──
    heading(doc, '1. System Overview', 1)
    para(doc,
         'ConceptGrade is a 5-stage automated short-answer grading system that grounds '
         'scoring in a domain Knowledge Graph (KG) rather than relying solely on LLM judgment.')

    doc.add_paragraph()
    heading(doc, '5-Stage Pipeline', 2)

    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(['Stage', 'Name', 'Description']):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    pipeline = [
        ('1', 'Question Parsing',    'Extract expected concepts from rubric — Gemini 2.5 Flash'),
        ('2', 'Concept Matching',    'Semantic match of student answer vs KG concepts (TF-IDF + embeddings)'),
        ('3a', 'LRM Verifier',       'Reasoning model verifies domain logic; exposes chain-of-thought trace'),
        ('3b', 'Trace Parser',       'Parses reasoning trace → SUPPORTS / CONTRADICTS / UNCERTAIN steps'),
        ('4', 'Chain Coverage',      '% of expected KG relationship chain covered by student concepts'),
        ('5', 'Score Aggregation',   'Weighted combination → final 0–5 score'),
    ]
    for row in pipeline:
        table_row(t, row)

    doc.add_paragraph()
    heading(doc, 'Conditions Compared', 2)
    bullets = [
        ('C_LLM',        'Pure LLM baseline (Gemini 2.5 Flash, no KG grounding)'),
        ('C5',           'Full 5-stage ConceptGrade pipeline (KG-grounded, no LRM trace adjustment)'),
        ('LRM-adjusted', 'C5 score ± scaled net_delta from the verifier\'s parsed reasoning trace'),
    ]
    for name, desc in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{name}: ').bold = True
        p.add_run(desc)

    doc.add_paragraph()
    para(doc, 'Research Question', bold=True)
    para(doc,
         'Does integrating a reasoning model\'s chain-of-thought trace (Stage 3b) as a '
         'post-hoc score adjustment improve grading accuracy beyond the 5-stage KG pipeline alone?',
         italic=True)

    doc.add_page_break()

    # ── 2. Datasets ──
    heading(doc, '2. Datasets', 1)

    t2 = doc.add_table(rows=1, cols=4)
    t2.style = 'Table Grid'
    for i, h in enumerate(['Dataset', 'Domain', 'N answers', 'Score range']):
        t2.rows[0].cells[i].text = h
        t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    table_row(t2, ['Mohler',      'CS: Data Structures',    '120',          '0–5'])
    table_row(t2, ['DigiKlausur', 'Neural Networks / DL',   '300 (sampled)','0–5'])

    # ── 3. LRM Verifier Details ──
    doc.add_paragraph()
    heading(doc, '3. LRM Verifier Details', 1)
    details = [
        'Primary model: DeepSeek-R1 (deepseek-reasoner) via DeepSeek API',
        'Fill-in model: Gemini 2.5 Flash with thinking (include_thoughts=True, budget=8192)',
        'Prompt inputs: domain, question, student answer, matched KG concepts, missing concepts, chain coverage %',
        'Trace parsing: each reasoning sentence → classification (SUPPORTS / CONTRADICTS / UNCERTAIN) + confidence_delta (±0.05–0.15)',
        'net_delta: sum of all step confidence_deltas per answer',
        'Score adjustment: clip(c5_score + scale × net_delta,  0, 5)',
    ]
    for d in details:
        doc.add_paragraph(d, style='List Bullet')

    # ── 4. Ablation Results ──
    doc.add_page_break()
    heading(doc, '4. Ablation Results', 1)

    for ds_key, ds_label in [('mohler', 'Mohler (n=120, CS domain)'),
                              ('digiklausur', 'DigiKlausur (n=300, Neural Networks)')]:
        heading(doc, f'4.x  {ds_label}', 2)
        s = summary[ds_key]

        t3 = doc.add_table(rows=1, cols=5)
        t3.style = 'Table Grid'
        for i, h in enumerate(['Metric', 'C_LLM', 'C5', 'LRM raw (scale=1)', 'LRM calibrated']):
            t3.rows[0].cells[i].text = h
            t3.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        def fmt(v):
            return f'{v:.4f}' if isinstance(v, float) else str(v)

        table_row(t3, ['MAE',
                        fmt(s['mae_cllm']),
                        fmt(s['mae_c5']),
                        fmt(s['mae_lrm_raw']),
                        fmt(s['mae_lrm_calibrated'])], bold=True)
        table_row(t3, ['vs C_LLM', '—',
                        f'{s["c5_vs_cllm_pct"]:+.1f}%', '—', '—'])
        table_row(t3, ['vs C5', '—', '—',
                        f'{s["lrm_raw_vs_c5_pct"]:+.1f}%',
                        f'{s["lrm_cal_vs_c5_pct"]:+.1f}%'])
        table_row(t3, ['Optimal scale', '—', '—', '1.0', fmt(s['lrm_scale'])])

        doc.add_paragraph()

    # ── 5. Review Questions ──
    doc.add_page_break()
    heading(doc, '5. Key Questions for Review', 1)

    questions = [
        (
            'Q1 — Why does LRM scale=0 optimise DigiKlausur?',
            'The calibration grid search finds optimal LRM scale = 0.0 for DigiKlausur, meaning any '
            'non-zero LRM adjustment worsens MAE. Three possible explanations:\n'
            '(a) Net_delta bias: the verifier assigns more CONTRADICTS than SUPPORTS (38% vs 34%), '
            'systematically pushing scores below the already-accurate C5 baseline.\n'
            '(b) Domain mismatch: DeepSeek-R1\'s internal model of "neural network concepts" may not '
            'match the specific course rubric — the KG (built from course material) is better calibrated.\n'
            '(c) Score ceiling effect: DigiKlausur C5 MAE is already 1.05. Adjustments that help on '
            'harder problems (Mohler MAE 2.25) may overshoot on an already-tight baseline.\n'
            'Which explanation is most plausible? What additional analysis would differentiate them?'
        ),
        (
            'Q2 — Why does C5 not improve over C_LLM on Mohler?',
            'On Mohler, C5 MAE equals C_LLM MAE (both 2.25). The LRM trace adjustment is the only '
            'improvement. Is this a KG quality problem (auto-generated KG, not expert-curated) or a '
            'fundamental limitation of concept-chain grading for CS data structures questions?'
        ),
        (
            'Q3 — Binary veto vs continuous delta',
            'Current use: LRM net_delta → continuous score adjustment. '
            'Alternative: LRM valid flag (72.5% Mohler, 58.7% DigiKlausur) as a binary veto '
            '(if valid=False → reduce score by fixed penalty). '
            'Would a binary validity gate outperform the continuous delta adjustment?'
        ),
        (
            'Q4 — Statistical significance',
            'Current analysis uses MAE point estimates only. With n=120 and n=300, is the 7.9% '
            'Mohler improvement statistically meaningful? Is Wilcoxon signed-rank on per-sample '
            'absolute errors the right test here?'
        ),
        (
            'Q5 — VIS framing of LRM trace visualisation',
            'For IEEE VIS 2027, the contribution must be the visualization + interaction design, not '
            'ML accuracy. How would you frame the LRM trace visualisation panel (SUPPORTS / '
            'CONTRADICTS / UNCERTAIN steps with KG node references and confidence deltas) as a novel '
            'VA contribution distinct from existing explainable AI dashboards?'
        ),
    ]

    for title, body in questions:
        para(doc, title, bold=True)
        doc.add_paragraph(body)
        doc.add_paragraph()

    # ── 6. Proposed Next Steps ──
    heading(doc, '6. Proposed Next Steps', 1)
    steps = [
        'Binary validity gate: implement valid=False → score × 0.85 and compare against delta adjustment',
        'Separate-model ablation: rerun 30 DigiKlausur samples with Gemini thinking only, compare traces to DeepSeek',
        'KG quality analysis: check Mohler KG node count per question; correlate with C5 improvement',
        'Statistical tests: Wilcoxon signed-rank on per-sample absolute errors for each condition pair',
        'User study: Condition A (summary card) vs Condition B (full VA dashboard with trace panel)',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    # ── Save ──
    out = RESULTS_DIR / 'conceptgrade_lrm_ablation_review.docx'
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    build_doc()
